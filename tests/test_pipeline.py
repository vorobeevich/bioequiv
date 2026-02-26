#!/usr/bin/env python3
"""
Тесты пайплайна: Stage 1, Stage 2 (все источники), data integrity.
Покрытие: Stage1, Vidal (препарат + вещество), e-Drug3D, OSP, CVintra/PMC,
          DrugBank, ОХЛП, fuzzy-reject, LLM (если ключ есть).

Запуск:
    python tests/test_pipeline.py        # встроенный runner
    python -m pytest tests/ -v           # через pytest
"""

import sys
import os
import csv

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pipeline.stage1 import find_original, find_all_by_inn
from pipeline.stage2 import find_pk_params
from pipeline.models import PK_PARAM_LABELS, PKParams

# ═══════════════════════════════════════════════
# Stage 1 — реестр ЕАЭС
# ═══════════════════════════════════════════════

def test_stage1_exact():
    """Точное совпадение — находим оригинальный препарат."""
    results = find_all_by_inn("амлодипин")
    assert len(results) >= 1
    originals = [d for d in results if d.drug_kind == "оригинальный"]
    assert len(originals) >= 1, "Должен быть оригинальный препарат"
    assert originals[0].match_type == "exact"
    assert "Норваск" in originals[0].trade_names

def test_stage1_fuzzy():
    """Fuzzy-поиск при опечатке (удвоенная буква)."""
    results = find_all_by_inn("амлодиппин")
    assert len(results) >= 1, "Fuzzy: амлодиппин должен найти амлодипин"
    assert results[0].match_type == "fuzzy"
    assert results[0].match_score >= 80

def test_stage1_not_found():
    """Несуществующий МНН — пустой список."""
    results = find_all_by_inn("абвгдежзиклмноп123")
    assert len(results) == 0

def test_stage1_find_original():
    """find_original возвращает оригинальный препарат."""
    drug = find_original("ибупрофен")
    assert drug is not None
    assert drug.drug_kind == "оригинальный"
    assert drug.matched_inn == "ибупрофен"

def test_stage1_multiple_kinds():
    """Для ибупрофена есть и оригиналы, и генерики."""
    results = find_all_by_inn("ибупрофен")
    kinds = {d.drug_kind for d in results}
    assert "оригинальный" in kinds
    assert len(kinds) >= 2

def test_stage1_atc_code():
    """Оригинальный препарат должен иметь АТХ-код."""
    drug = find_original("амлодипин")
    assert drug is not None
    assert drug.atc_code.startswith("C"), f"Ожидался кардио АТХ, получили: {drug.atc_code}"

def test_stage1_latin_name():
    """trade_names не пустые."""
    drug = find_original("метформин")
    assert drug is not None
    assert len(drug.trade_names) > 0

def test_stage1_sertraline():
    """Сертралин — Золофт."""
    results = find_all_by_inn("сертралин")
    originals = [d for d in results if d.drug_kind == "оригинальный"]
    assert len(originals) >= 1
    assert "Золофт" in originals[0].trade_names


# ═══════════════════════════════════════════════
# Vidal: поиск вещества
# ═══════════════════════════════════════════════

def test_vidal_molecule_exact():
    """Точный поиск вещества."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_molecule("ибупрофен")
    assert r is not None
    assert r["match_type"] in ("exact", "exact_latin")
    assert r.get("name_latin", "").lower() == "ibuprofen"

def test_vidal_molecule_url():
    """Молекула должна иметь URL."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_molecule("ибупрофен")
    assert r is not None
    assert r.get("url", "").startswith("https://")

def test_vidal_molecule_has_pk_text():
    """Для популярных веществ должен быть ФК текст."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_molecule("амлодипин")
    assert r is not None
    pk = r.get("pharmacokinetics", "")
    assert len(pk) > 50, "Ожидается непустой текст ФК"

def test_vidal_molecule_amlodipine_latin():
    """Мост рус→лат для амлодипина."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_molecule("амлодипин")
    assert r is not None
    assert r.get("name_latin", "").lower() == "amlodipine"

def test_vidal_molecule_not_found():
    """Несуществующее вещество — None."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_molecule("абвгдеж12345")
    assert r is None

def test_vidal_molecule_fuzzy_validated_reject():
    """Совсем разные слова не должны матчиться."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_molecule("zoloft1234xyz")
    assert r is None


# ═══════════════════════════════════════════════
# Vidal: поиск препарата
# ═══════════════════════════════════════════════

def test_vidal_drug_exact():
    """Точный поиск препарата по торговому названию."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_drug("Нурофен")
    if r:
        assert r.get("level") == "drug"
        assert "ibuprofen" in r.get("name_latin", "").lower() or \
               "ибупрофен" in r.get("molecule_ru", "").lower()

def test_vidal_drug_sertraline_in_pipeline():
    """Anti-regression: сертралин через пайплайн не должен получить препарат с Золототысячником."""
    drug = find_original("сертралин")
    assert drug is not None
    res = find_pk_params(drug, use_llm=False)
    # Если vidal_drug_result заполнен — его молекула должна быть сертралином
    if res.vidal_drug_result:
        mol = res.vidal_drug_result.get("molecule_ru", "").lower()
        name_lat = res.vidal_drug_result.get("name_latin", "").lower()
        assert "сертралин" in mol or "sertralin" in name_lat, \
            f"Stage2 принял ложный матч Видаль: {res.vidal_drug_result.get('drug_name')} → {mol}"

def test_vidal_drug_has_level_field():
    """Результат search_drug должен содержать level=drug."""
    from pipeline.stage2_sources import vidal
    r = vidal.search_drug("Нурофен")
    if r:
        assert r.get("level") == "drug"


# ═══════════════════════════════════════════════
# e-Drug3D
# ═══════════════════════════════════════════════

def test_edrug3d_amlodipine():
    """Амлодипин — T½ должен быть около 40 ч."""
    from pipeline.stage2_sources import edrug3d
    r = edrug3d.search("amlodipine")
    assert r is not None
    params = r.get("params", {})
    t_half = params.get("t_half_h")
    assert t_half is not None
    assert 30 <= t_half.value <= 50, f"T½ амлодипина ожидается ~40 ч, получили {t_half.value}"

def test_edrug3d_ibuprofen_t_half():
    """Ибупрофен — T½ ~2 ч."""
    from pipeline.stage2_sources import edrug3d
    r = edrug3d.search("ibuprofen")
    assert r is not None
    t_half = r.get("params", {}).get("t_half_h")
    assert t_half is not None
    assert 1 <= t_half.value <= 4

def test_edrug3d_not_found():
    """Несуществующий препарат — None."""
    from pipeline.stage2_sources import edrug3d
    r = edrug3d.search("абвгдеж12345xyz")
    assert r is None

def test_edrug3d_match_type_field():
    """Результат должен содержать match_type."""
    from pipeline.stage2_sources import edrug3d
    r = edrug3d.search("amlodipine")
    assert r is not None
    assert "match_type" in r

def test_edrug3d_params_are_pkvalue():
    """Параметры должны быть PKValue объектами."""
    from pipeline.stage2_sources import edrug3d
    from pipeline.models import PKValue
    r = edrug3d.search("amlodipine")
    assert r is not None
    for pname, pval in r.get("params", {}).items():
        assert isinstance(pval, PKValue), f"{pname} должен быть PKValue"
        assert pval.value is not None
        assert pval.unit != ""


# ═══════════════════════════════════════════════
# OSP
# ═══════════════════════════════════════════════

def test_osp_verapamil_cvintra():
    """Верапамил есть в OSP с CVintra."""
    from pipeline.stage2_sources import osp
    r = osp.search("verapamil")
    assert r is not None
    cv = r.get("params", {}).get("cvintra_pct")
    assert cv is not None, "Верапамил должен иметь CVintra в OSP"
    assert 0 < cv.value < 100

def test_osp_midazolam():
    """Мидазолам есть в OSP."""
    from pipeline.stage2_sources import osp
    r = osp.search("midazolam")
    assert r is not None
    assert len(r.get("params", {})) >= 1

def test_osp_not_found():
    """Несуществующее вещество — None."""
    from pipeline.stage2_sources import osp
    r = osp.search("абвгдеж12345xyz")
    assert r is None

def test_osp_cvintra_value_range():
    """CVintra из OSP должен быть в разумных пределах."""
    from pipeline.stage2_sources import osp
    r = osp.search("verapamil")
    if r:
        cv = r.get("params", {}).get("cvintra_pct")
        if cv:
            assert 1 < cv.value < 200, f"CVintra вне диапазона: {cv.value}"


# ═══════════════════════════════════════════════
# CVintra PMC
# ═══════════════════════════════════════════════

def test_cvintra_pmc_amlodipine():
    """Амлодипин есть в PMC базе."""
    from pipeline.stage2_sources import cvintra_pmc
    r = cvintra_pmc.search("amlodipine")
    assert r is not None
    assert r.get("cvintra_cmax_pct") is not None

def test_cvintra_pmc_atorvastatin():
    """Аторвастатин — HVD (CVintra >30%)."""
    from pipeline.stage2_sources import cvintra_pmc
    r = cvintra_pmc.search("atorvastatin")
    assert r is not None
    cv = r.get("cvintra_cmax_pct")
    assert cv is not None
    assert cv >= 30, f"Аторвастатин должен быть HVD (>=30%), получили {cv}"

def test_cvintra_pmc_has_sample_size():
    """PMC запись должна содержать рекомендацию по выборке."""
    from pipeline.stage2_sources import cvintra_pmc
    r = cvintra_pmc.search("amlodipine")
    assert r is not None
    assert r.get("sample_size_80pwr") or r.get("sample_size_90pwr"), \
        "Должна быть рекомендация по размеру выборки"

def test_cvintra_pmc_not_found():
    """Неизвестное вещество — None."""
    from pipeline.stage2_sources import cvintra_pmc
    r = cvintra_pmc.search("абвгдеж12345xyz")
    assert r is None

def test_cvintra_pmc_params_pkvalue():
    """cvintra_pct должен быть PKValue."""
    from pipeline.stage2_sources import cvintra_pmc
    from pipeline.models import PKValue
    r = cvintra_pmc.search("amlodipine")
    assert r is not None
    cv = r.get("params", {}).get("cvintra_pct")
    assert cv is not None
    assert isinstance(cv, PKValue)
    assert cv.unit == "%"


# ═══════════════════════════════════════════════
# DrugBank
# ═══════════════════════════════════════════════

def test_drugbank_ibuprofen():
    """Ибупрофен есть в DrugBank."""
    from pipeline.stage2_sources import drugbank
    r = drugbank.search("ibuprofen")
    assert r is not None

def test_drugbank_url():
    """DrugBank: результат должен содержать URL."""
    from pipeline.stage2_sources import drugbank
    r = drugbank.search("ibuprofen")
    assert r is not None
    assert "drugbank" in r.get("url", "").lower()

def test_drugbank_has_halflife():
    """DrugBank: должно быть поле half_life."""
    from pipeline.stage2_sources import drugbank
    r = drugbank.search("ibuprofen")
    assert r is not None
    assert len(r.get("half_life", "")) > 0

def test_drugbank_not_found():
    """Несуществующее вещество — None."""
    from pipeline.stage2_sources import drugbank
    r = drugbank.search("абвгдеж12345xyz")
    assert r is None

def test_drugbank_amlodipine():
    """Амлодипин есть в DrugBank."""
    from pipeline.stage2_sources import drugbank
    r = drugbank.search("amlodipine")
    assert r is not None
    assert r.get("drugbank_id", "").startswith("DB")


# ═══════════════════════════════════════════════
# ОХЛП
# ═══════════════════════════════════════════════

def test_ohlp_by_inn():
    """ОХЛП: поиск по МНН."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    r = ohlp.search("ибупрофен")
    assert r is not None
    assert r.get("level") == "substance"
    assert len(r.get("pk_text", "")) > 50

def test_ohlp_by_trade_name():
    """ОХЛП: поиск по торговому названию имеет приоритет."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    r = ohlp.search("ибупрофен", trade_name="Бумидол")
    assert r is not None
    # Если нашло по торговому — уровень drug, иначе substance (оба допустимы)
    assert r.get("level") in ("drug", "substance")

def test_ohlp_level_drug_preferred():
    """ОХЛП: при наличии trade_name уровень drug предпочтительнее."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    r_inn = ohlp.search("ибупрофен")
    r_tn = ohlp.search("ибупрофен", trade_name="Бумидол")
    assert r_inn is not None
    # Если r_tn нашёл по торговому — его уровень drug
    if r_tn and r_tn.get("level") == "drug":
        assert r_tn.get("matched_trade_name", "").lower() == "бумидол"

def test_ohlp_has_sections():
    """ОХЛП: должны быть основные разделы."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    r = ohlp.search("ибупрофен")
    assert r is not None
    # Хотя бы pk_text или dosing_text должны быть непустыми
    has_text = any(len(r.get(f, "")) > 30
                   for f in ("pk_text", "dosing_text", "contra_text"))
    assert has_text

def test_ohlp_result_has_level_field():
    """Результат ОХЛП должен содержать поле level."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    r = ohlp.search("ибупрофен")
    assert r is not None
    assert "level" in r
    assert r["level"] in ("drug", "substance")

def test_ohlp_not_found():
    """ОХЛП: несуществующий препарат — None."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    r = ohlp.search("абвгдеж12345", trade_name="абвгдеж12345")
    assert r is None


# ═══════════════════════════════════════════════
# Stage 2 полный пайплайн (без LLM)
# ═══════════════════════════════════════════════

def test_stage2_log_structure():
    """Лог должен содержать метки всех ключевых шагов."""
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=False)
    log = "\n".join(res.log)
    for marker in ("[2.0]", "[2.1]", "[2.2]", "Итого:"):
        assert marker in log, f"В логе нет метки: {marker}"

def test_stage2_amlodipine_t_half():
    """Амлодипин: T½ ~40 ч из e-Drug3D."""
    drug = find_original("амлодипин")
    res = find_pk_params(drug, use_llm=False)
    assert res.edrug3d_result is not None
    t_half = res.pk.t_half_h
    assert t_half is not None
    assert 30 <= t_half.value <= 50

def test_stage2_ibuprofen_sources():
    """Ибупрофен: находим Видаль, e-Drug3D, DrugBank."""
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=False)
    assert res.vidal_mol_result is not None, "Видаль не нашёл вещество"
    assert res.edrug3d_result is not None, "e-Drug3D не нашёл вещество"
    assert res.drugbank_result is not None, "DrugBank не нашёл вещество"
    assert res.name_latin.lower() == "ibuprofen"

def test_stage2_ibuprofen_t_half():
    """Ибупрофен: T½ ~2 ч."""
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=False)
    t_half = res.pk.t_half_h
    assert t_half is not None
    assert 1 <= t_half.value <= 4

def test_stage2_amlodipine_cvintra_pmc():
    """Амлодипин: CVintra есть в PMC базе."""
    drug = find_original("амлодипин")
    res = find_pk_params(drug, use_llm=False)
    assert res.cvintra_pmc_result is not None
    cv = res.cvintra_pmc_result.get("cvintra_cmax_pct")
    assert cv is not None

def test_stage2_ohlp_in_pipeline():
    """ОХЛП подключается в пайплайне."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=False)
    assert res.ohlp_result is not None
    assert "ОХЛП" in "\n".join(res.log)

def test_stage2_ohlp_level_logged():
    """Уровень ОХЛП (drug/substance) должен быть в логе."""
    from pipeline.stage2_sources import ohlp
    if not ohlp.OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=False)
    log = "\n".join(res.log)
    assert "уровень:" in log, "Уровень ОХЛП должен быть в логе"

def test_stage2_pkparams_is_pkparams():
    """Результат Stage 2 должен иметь корректный PKParams объект."""
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=False)
    assert isinstance(res.pk, PKParams)

def test_stage2_vildagliptin():
    """Вилдаглиптин: Видаль и DrugBank находят что-то."""
    drug = find_original("вилдаглиптин")
    res = find_pk_params(drug, use_llm=False)
    assert res.name_latin.lower() == "vildagliptin"
    assert res.drugbank_result is not None

def test_stage2_metformin():
    """Метформин: полный пайплайн без ошибок."""
    drug = find_original("метформин")
    assert drug is not None
    res = find_pk_params(drug, use_llm=False)
    assert res.vidal_mol_result is not None

def test_stage2_sertraline_no_wrong_vidal():
    """Сертралин: Видаль не должен возвращать ложный fuzzy без LLM валидации."""
    drug = find_original("сертралин")
    assert drug is not None
    res = find_pk_params(drug, use_llm=False)
    # Если vidal_drug_result есть — оно должно содержать «сертралин» или «sertraline»
    if res.vidal_drug_result:
        mol = res.vidal_drug_result.get("molecule_ru", "").lower()
        name_lat = res.vidal_drug_result.get("name_latin", "").lower()
        assert "сертралин" in mol or "sertralin" in name_lat, \
            f"Видаль вернул неверный препарат: {res.vidal_drug_result.get('drug_name')}"

def test_stage2_no_cmax_molar_ibuprofen():
    """Ибупрофен: Cmax в молярных единицах не должен подставляться в pk.cmax."""
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=False)
    cmax = res.pk.cmax
    if cmax is not None:
        assert cmax.unit in ("нг/мл", "ng/mL"), \
            f"Единица Cmax неожиданная: {cmax.unit}"


# ═══════════════════════════════════════════════
# Data integrity
# ═══════════════════════════════════════════════

def test_data_eaeu_registry():
    """Реестр ЕАЭС: есть данные."""
    from pipeline.config import EAEU_REGISTRY_CSV
    with open(EAEU_REGISTRY_CSV, encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    assert len(rows) >= 1000, f"Слишком мало записей в реестре: {len(rows)}"

def test_data_vidal_molecules():
    """Видаль вещества: есть данные с латиницей."""
    from pipeline.config import VIDAL_MOLECULES_CSV
    with open(VIDAL_MOLECULES_CSV, encoding="utf-8", errors="replace") as f:
        rows = list(csv.DictReader(f))
    assert len(rows) >= 1000
    has_latin = sum(1 for r in rows if r.get("name_latin", "").strip())
    assert has_latin >= 500, f"Мало веществ с латинским именем: {has_latin}"

def test_data_edrug3d():
    """e-Drug3D: файл есть и непустой."""
    from pipeline.config import EDRUG3D_CSV
    with open(EDRUG3D_CSV, encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    assert len(rows) >= 500

def test_data_osp():
    """OSP: файл есть и непустой."""
    from pipeline.config import OSP_CSV
    with open(OSP_CSV, encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    assert len(rows) >= 100

def test_data_cvintra_pmc():
    """CVintra PMC: 50+ веществ."""
    from pipeline.config import CVINTRA_PMC_CSV
    with open(CVINTRA_PMC_CSV, encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    assert len(rows) >= 50, f"CVintra PMC: мало записей: {len(rows)}"

def test_data_drugbank():
    """DrugBank: файл есть и непустой."""
    from pipeline.config import DRUGBANK_CSV
    with open(DRUGBANK_CSV, encoding="utf-8", errors="replace") as f:
        rows = list(csv.DictReader(f))
    assert len(rows) >= 1000

def test_data_ohlp():
    """ОХЛП: если включён — 1000+ записей."""
    from pipeline.config import OHLP_CSV, OHLP_ENABLED
    if not OHLP_ENABLED:
        print("  SKIP: OHLP CSV не найден")
        return
    with open(OHLP_CSV, encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        count = sum(1 for _ in reader)
    assert count >= 1000, f"ОХЛП: мало записей: {count}"


# ═══════════════════════════════════════════════
# FDA PSG
# ═══════════════════════════════════════════════

def test_fda_psg_file_exists():
    """fda_psg_parsed.csv должен существовать."""
    from pipeline.config import FDA_PSG_CSV
    assert os.path.exists(FDA_PSG_CSV), f"Файл не найден: {FDA_PSG_CSV}"

def test_fda_psg_row_count():
    """Должно быть 1000+ записей."""
    from pipeline.config import FDA_PSG_CSV, FDA_PSG_ENABLED
    if not FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    with open(FDA_PSG_CSV, encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    assert len(rows) >= 1000, f"Мало записей: {len(rows)}"

def test_fda_psg_search_amlodipine():
    """Амлодипин есть в FDA PSG (через Aliskiren+Amlodipine)."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("amlodipine")
    assert r is not None, "Амлодипин не найден в FDA PSG"
    assert r.get("source") == "fda_psg"

def test_fda_psg_carbamazepine_nti():
    """Карбамазепин — NTI препарат в FDA PSG."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("carbamazepine")
    assert r is not None
    assert r.get("is_nti"), "Карбамазепин должен быть помечен как NTI"

def test_fda_psg_carbamazepine_replicated():
    """Карбамазепин требует реплицированного дизайна."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("carbamazepine")
    assert r is not None
    assert r.get("is_replicated"), "Карбамазепин должен иметь replicated design"

def test_fda_psg_acamprosate_hvd():
    """Акампросат — HVD препарат."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("acamprosate")
    assert r is not None
    assert r.get("is_hvd"), "Акампросат должен быть помечен как HVD"

def test_fda_psg_cvintra_threshold():
    """CVintra threshold (30) для HVD препарата."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("acamprosate")
    assert r is not None
    assert r.get("cvintra_threshold") == 30, \
        f"CVintra порог акампросата = {r.get('cvintra_threshold')}, ожидается 30"

def test_fda_psg_has_design():
    """Результат должен содержать дизайн исследования."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("metformin")
    assert r is not None
    assert len(r.get("design_fasting", "")) > 10, "Дизайн натощак должен быть непустым"

def test_fda_psg_has_pdf_url():
    """Результат должен содержать ссылку на PDF."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("metformin")
    assert r is not None
    assert r.get("pdf_url", "").startswith("https://"), \
        f"Неверный URL: {r.get('pdf_url')}"

def test_fda_psg_not_found():
    """Несуществующее вещество → None."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("абвгдеж12345xyz")
    assert r is None

def test_fda_psg_result_fields():
    """Результат должен содержать все обязательные поля."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    r = fda_psg.search("metformin")
    assert r is not None
    for field in ("source", "substance", "form_route", "match_type", "match_score",
                  "is_replicated", "is_hvd", "is_nti"):
        assert field in r, f"Поле '{field}' отсутствует в результате"

def test_fda_psg_search_all():
    """search_all возвращает несколько записей для popular drug."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    results = fda_psg.search_all("metformin")
    assert len(results) >= 2, f"Ожидалось 2+ записей метформина, получили {len(results)}"

def test_fda_psg_in_pipeline():
    """FDA PSG интегрирован в Stage 2 пайплайн."""
    from pipeline.stage2_sources import fda_psg
    if not fda_psg.FDA_PSG_ENABLED:
        print("  SKIP: fda_psg_parsed.csv не найден")
        return
    drug = find_original("метформин")
    assert drug is not None
    res = find_pk_params(drug, use_llm=False)
    assert res.fda_psg_result is not None, "FDA PSG должен быть в результате Stage 2"
    log = "\n".join(res.log)
    assert "FDA PSG" in log


# ═══════════════════════════════════════════════
# Stage 2 с LLM (только если ключ задан)
# ═══════════════════════════════════════════════

def test_stage2_ibuprofen_with_llm():
    """Ибупрофен с LLM: не менее 2 параметров."""
    from pipeline.config import DEEPSEEK_API_KEY
    if not DEEPSEEK_API_KEY:
        print("  SKIP: DEEPSEEK_API_KEY не задан")
        return
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=True)
    pk = res.pk
    filled = pk.filled_params()
    assert len(filled) >= 2, f"LLM: ожидалось 2+ параметра, получили {len(filled)}"
    assert pk.t_half_h is not None

def test_stage2_amlodipine_with_llm():
    """Амлодипин с LLM: не менее 3 параметров."""
    from pipeline.config import DEEPSEEK_API_KEY
    if not DEEPSEEK_API_KEY:
        print("  SKIP: DEEPSEEK_API_KEY не задан")
        return
    drug = find_original("амлодипин")
    res = find_pk_params(drug, use_llm=True)
    filled = res.pk.filled_params()
    assert len(filled) >= 3, f"LLM: ожидалось 3+ параметра, получили {len(filled)}"

def test_stage2_source_tag_in_pkvalue():
    """LLM: source_tag в PKValue должен содержать имя источника."""
    from pipeline.config import DEEPSEEK_API_KEY
    if not DEEPSEEK_API_KEY:
        print("  SKIP: DEEPSEEK_API_KEY не задан")
        return
    drug = find_original("ибупрофен")
    res = find_pk_params(drug, use_llm=True)
    for pname, pval in res.pk.filled_params().items():
        assert pval.source, f"{pname}: source пустой"


# ═══════════════════════════════════════════════
# Stage 3 — sample size, timepoints, synopsis
# ═══════════════════════════════════════════════

def test_sample_size_basic():
    """Sample size: CV 25%, power 80% → разумное n."""
    from pipeline.sample_size import calc_sample_size
    r = calc_sample_size(25.0, power=0.80)
    assert r["n_evaluable"] >= 12
    assert r["n_total"] >= r["n_evaluable"]
    assert r["n_total"] % 2 == 0

def test_sample_size_nti():
    """Sample size: NTI (theta 1.1111) → больше n чем стандарт при том же CV."""
    from pipeline.sample_size import calc_sample_size
    r_std = calc_sample_size(20.0, power=0.80, theta=1.25)
    r_nti = calc_sample_size(20.0, power=0.80, theta=1.1111)
    assert r_nti["n_evaluable"] > r_std["n_evaluable"]

def test_sample_size_high_cv():
    """Sample size: высокий CV → больше n."""
    from pipeline.sample_size import calc_sample_size
    r_low = calc_sample_size(15.0, power=0.80)
    r_high = calc_sample_size(45.0, power=0.80)
    assert r_high["n_evaluable"] > r_low["n_evaluable"]

def test_sample_size_replicated():
    """Replicated design: n_evaluable уменьшается (0.75x)."""
    from pipeline.sample_size import calc_sample_size
    r_2x2 = calc_sample_size(35.0, power=0.80, design="2x2")
    r_rep = calc_sample_size(35.0, power=0.80, design="replicated")
    assert r_rep["n_evaluable"] <= r_2x2["n_evaluable"]

def test_sample_size_formula_note():
    """Formula note содержит ключевые слова."""
    from pipeline.sample_size import calc_sample_size
    r = calc_sample_size(25.0)
    note = r["formula_note"]
    assert "CVintra" in note
    assert "мощность" in note
    assert "Решение" in note

def test_determine_design_standard():
    """Дизайн: CV < 30%, не NTI, не HVD → стандартный 2x2."""
    from pipeline.sample_size import determine_design
    d = determine_design(20.0)
    assert d["design"] == "2x2"
    assert d["theta"] == 1.25

def test_determine_design_hvd():
    """Дизайн: HVD → replicated."""
    from pipeline.sample_size import determine_design
    d = determine_design(35.0, is_hvd=True)
    assert d["design"] == "replicated"

def test_determine_design_nti():
    """Дизайн: NTI → theta 1.1111."""
    from pipeline.sample_size import determine_design
    d = determine_design(15.0, is_nti=True)
    assert d["theta"] == 1.1111
    assert "90.00–111.11%" in d["be_limits"]

def test_timepoints_basic():
    """Timepoints: разумное кол-во точек для ибупрофена."""
    from pipeline.timepoints import generate_timepoints
    r = generate_timepoints(1.5, 2.0)
    assert 12 <= r["n_samples"] <= 25
    assert r["timepoints_h"][0] == 0.0
    assert r["end_time_h"] <= 72

def test_timepoints_long_thalf():
    """Timepoints: длинный T½ → отбор до 72 ч."""
    from pipeline.timepoints import generate_timepoints
    r = generate_timepoints(8.0, 40.0)
    assert r["end_time_h"] == 72.0
    assert r["n_samples"] >= 15

def test_timepoints_blood_volume():
    """Timepoints: общий объём крови рассчитывается."""
    from pipeline.timepoints import generate_timepoints
    r = generate_timepoints(2.0, 6.0)
    assert r["total_blood_2periods_ml"] > 0
    assert r["blood_per_sample_ml"] == 5.0

def test_timepoints_schedule_text():
    """Timepoints: schedule_text содержит '0 (до приёма)'."""
    from pipeline.timepoints import generate_timepoints
    r = generate_timepoints(2.0, 6.0)
    assert "0 (до приёма)" in r["schedule_text"]

def test_stage3_compute_derived():
    """Stage 3 compute_derived: подставить значения из Stage 2."""
    from pipeline.stage3 import Stage3Input, compute_derived
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    drug_info = DrugInfo(query_inn="тест", matched_inn="тест")
    s2 = Stage2Result()
    s2.pk = PKParams(
        t_half_h=PKValue(value=6.0, unit="ч"),
        tmax_h=PKValue(value=2.0, unit="ч"),
        cvintra_pct=PKValue(value=25.0, unit="%"),
    )

    inp = Stage3Input(drug_info=drug_info, s2=s2, test_drug_name="Test")
    computed = compute_derived(inp)
    assert computed["t_half"] == 6.0
    assert computed["tmax"] == 2.0
    assert computed["cv_intra"] == 25.0
    assert computed["washout_days"] == 2  # ceil(5*6/24) = 2
    assert computed["vomit_criterion_h"] == 4.0
    assert computed["sample_size"] is not None
    assert computed["timepoints"] is not None

def test_stage3_generate_docx():
    """Stage 3 generate_docx: возвращает непустые байты."""
    from pipeline.stage3 import Stage3Input, Stage3Result, generate_docx

    result = Stage3Result(
        synopsis={
            "protocol_title": "Тест протокол",
            "phase": "Исследование биоэквивалентности",
            "inn": "тестостерон",
            "test_drug_name": "Генерик",
        },
        sources_used=[{"name": "Видаль", "url": "https://www.vidal.ru", "type": "drug"}],
    )
    docx_bytes = generate_docx(result)
    assert len(docx_bytes) > 1000
    assert docx_bytes[:4] == b"PK\x03\x04"  # ZIP magic (docx is a zip)

def test_stage3_synopsis_no_llm():
    """Stage 3 generate_synopsis без LLM: только расчётные поля."""
    from pipeline.stage3 import Stage3Input, generate_synopsis
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    drug_info = DrugInfo(query_inn="амлодипин", matched_inn="амлодипин", trade_names="Норваск")
    s2 = Stage2Result()
    s2.pk = PKParams(
        t_half_h=PKValue(value=40.0, unit="ч"),
        tmax_h=PKValue(value=8.0, unit="ч"),
        cvintra_pct=PKValue(value=25.0, unit="%"),
    )

    inp = Stage3Input(drug_info=drug_info, s2=s2, test_drug_name="Амлодипин-Тест", sponsor="ООО Тест")
    result = generate_synopsis(inp, llm_fn=None)

    assert result.synopsis["inn"] == "амлодипин"
    assert result.synopsis["test_drug_name"] == "Амлодипин-Тест"
    assert result.synopsis["reference_drug_name"] == "Норваск"
    assert result.synopsis["sponsor"] == "ООО Тест"
    assert "n_total" in result.synopsis
    assert int(result.synopsis["n_total"]) >= 12
    assert result.sources_used


def test_stage3_parse_llm_json_clean():
    """_parse_llm_json: чистый JSON."""
    from pipeline.stage3 import _parse_llm_json
    r = _parse_llm_json('{"a": "hello", "b": "world"}')
    assert r == {"a": "hello", "b": "world"}

def test_stage3_parse_llm_json_markdown():
    """_parse_llm_json: JSON обёрнутый в markdown-блок."""
    from pipeline.stage3 import _parse_llm_json
    r = _parse_llm_json('```json\n{"key": "value"}\n```')
    assert r == {"key": "value"}

def test_stage3_parse_llm_json_with_preamble():
    """_parse_llm_json: JSON с текстом до и после."""
    from pipeline.stage3 import _parse_llm_json
    r = _parse_llm_json('Here is your answer:\n{"x": 1}\nDone!')
    assert r == {"x": 1}

def test_stage3_parse_llm_json_invalid():
    """_parse_llm_json: невалидный JSON → пустой dict."""
    from pipeline.stage3 import _parse_llm_json
    r = _parse_llm_json('not json at all')
    assert r == {}

def test_stage3_collect_all_data_keys():
    """collect_all_data: возвращает правильные ключи для полных данных."""
    from pipeline.stage3 import Stage3Input, collect_all_data
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo

    drug_info = DrugInfo(query_inn="тест", matched_inn="тест")
    s2 = Stage2Result()
    s2.vidal_drug_result = {"pharmacokinetics": "Всасывание быстрое...", "drug_name": "Тест"}
    s2.vidal_mol_result = {"pharmacokinetics": "PK text", "contraindications": "Аллергия", "name_ru": "тест"}
    s2.ohlp_result = {
        "pk_text": "Фармакокинетика длинный текст 123",
        "contra_text": "Противопоказания длинный текст 1",
        "adverse_text": "Побочные длинный текст тест 12345",
        "dosing_text": "Дозирование длинный текст тест 1",
    }
    s2.drugbank_result = {"absorption": "Rapidly absorbed", "half_life": "2 hours", "matched_name": "Test"}
    s2.fda_psg_result = {"design_fasting": "crossover", "analytes": "Ibuprofen", "strength": "400 mg"}

    inp = Stage3Input(drug_info=drug_info, s2=s2)
    data = collect_all_data(inp)

    assert "vidal_drug" in data
    assert "vidal_mol_pharmacokinetics" in data
    assert "vidal_mol_contraindications" in data
    assert "ohlp_pk_text" in data
    assert "ohlp_contra_text" in data
    assert "ohlp_adverse_text" in data
    assert "drugbank_absorption" in data
    assert "drugbank_half_life" in data
    assert "fda_psg_design_fasting" in data
    assert "fda_psg_analytes" in data

def test_stage3_collect_all_data_empty():
    """collect_all_data: пустой Stage2 → пустой dict."""
    from pipeline.stage3 import Stage3Input, collect_all_data
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo

    s2 = Stage2Result()
    inp = Stage3Input(drug_info=DrugInfo(query_inn="x", matched_inn="x"), s2=s2)
    data = collect_all_data(inp)
    assert data == {}

def test_stage3_llm_calls_mapping():
    """LLM_CALLS: 3 вызова определены с правильными полями."""
    from pipeline.stage3 import LLM_CALLS
    assert len(LLM_CALLS) == 3

    all_fields = []
    ids = set()
    for call in LLM_CALLS:
        assert "id" in call
        assert "fields" in call
        assert "data_keys" in call
        assert len(call["fields"]) >= 2
        ids.add(call["id"])
        all_fields.extend(call["fields"])

    assert ids == {"study_design_analysis", "criteria", "drug_safety"}

    expected = {
        "tasks", "study_design",
        "inclusion_criteria", "exclusion_criteria", "withdrawal_criteria",
        "test_drug_details", "reference_drug_details", "safety_analysis",
    }
    assert set(all_fields) == expected, f"Разница: {set(all_fields) ^ expected}"

def test_stage3_prompt_builds_without_error():
    """Промпты для всех 3 вызовов формируются без ошибок."""
    from pipeline.stage3 import (
        Stage3Input, compute_derived, collect_all_data,
        LLM_CALLS, _load_rule85, _build_prompt_for_call,
    )
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    drug_info = DrugInfo(query_inn="ибупрофен", matched_inn="ибупрофен", trade_names="Нурофен")
    s2 = Stage2Result()
    s2.pk = PKParams(
        t_half_h=PKValue(value=2.0, unit="ч"),
        tmax_h=PKValue(value=1.5, unit="ч"),
        cvintra_pct=PKValue(value=35.0, unit="%"),
    )
    s2.ohlp_result = {
        "contra_text": "Гиперчувствительность к ибупрофену 12345",
        "adverse_text": "Тошнота, рвота, головная боль и тд 12",
    }
    s2.fda_psg_result = {"design_fasting": "crossover", "analytes": "Ibuprofen"}

    inp = Stage3Input(drug_info=drug_info, s2=s2, test_drug_name="Ибупрофен-Тест", strength="400 мг")
    computed = compute_derived(inp)
    all_data = collect_all_data(inp)
    rule85 = _load_rule85()

    for call_def in LLM_CALLS:
        prompt = _build_prompt_for_call(call_def, inp, computed, all_data, rule85)
        assert len(prompt) > 200, f"{call_def['id']}: промпт слишком короткий"
        assert "ибупрофен" in prompt.lower(), f"{call_def['id']}: нет МНН в промпте"
        for fld in call_def["fields"]:
            assert fld in prompt, f"{call_def['id']}: поле {fld} не упомянуто в промпте"

def test_stage3_prompt_contains_relevant_data():
    """Промпт criteria содержит данные из ОХЛП (противопоказания)."""
    from pipeline.stage3 import (
        Stage3Input, compute_derived, collect_all_data,
        LLM_CALLS, _load_rule85, _build_prompt_for_call,
    )
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    drug_info = DrugInfo(query_inn="тест", matched_inn="тест", trade_names="ТестДраг")
    s2 = Stage2Result()
    s2.pk = PKParams(t_half_h=PKValue(value=5.0, unit="ч"), tmax_h=PKValue(value=2.0, unit="ч"))
    s2.ohlp_result = {
        "contra_text": "УНИКАЛЬНОЕ_ПРОТИВОПОКАЗАНИЕ_12345_ТЕСТ",
        "adverse_text": "УНИКАЛЬНОЕ_ПОБОЧНОЕ_ДЕЙСТВИЕ_67890_ТСТ",
    }

    inp = Stage3Input(drug_info=drug_info, s2=s2, test_drug_name="Тест-Генерик")
    computed = compute_derived(inp)
    all_data = collect_all_data(inp)
    rule85 = _load_rule85()

    criteria_call = [c for c in LLM_CALLS if c["id"] == "criteria"][0]
    prompt = _build_prompt_for_call(criteria_call, inp, computed, all_data, rule85)
    assert "УНИКАЛЬНОЕ_ПРОТИВОПОКАЗАНИЕ_12345_ТЕСТ" in prompt
    assert "УНИКАЛЬНОЕ_ПОБОЧНОЕ_ДЕЙСТВИЕ_67890_ТСТ" in prompt

    safety_call = [c for c in LLM_CALLS if c["id"] == "drug_safety"][0]
    prompt_safety = _build_prompt_for_call(safety_call, inp, computed, all_data, rule85)
    assert "УНИКАЛЬНОЕ_ПРОТИВОПОКАЗАНИЕ_12345_ТЕСТ" not in prompt_safety

def test_stage3_extract_rule85_sections():
    """_extract_rule85_sections: находит релевантные секции."""
    from pipeline.stage3 import _extract_rule85_sections, _load_rule85
    rule85 = _load_rule85()
    if not rule85:
        print("  SKIP: rule85_context.md не найден")
        return

    design_sections = _extract_rule85_sections(rule85, ["Дизайн исследования", "Проведение"])
    assert "Стандартный дизайн" in design_sections or "дизайн" in design_sections.lower()

    pk_sections = _extract_rule85_sections(rule85, ["ФК параметры", "Критерии БЭ"])
    assert "AUC" in pk_sections or "Cmax" in pk_sections

    empty = _extract_rule85_sections(rule85, ["НЕСУЩЕСТВУЮЩАЯ_СЕКЦИЯ_XYZ"])
    assert empty == ""

def test_stage3_synopsis_no_llm_missing_cv():
    """Stage 3 без LLM и без CVintra: не падает, ставит минимум."""
    from pipeline.stage3 import Stage3Input, generate_synopsis
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    drug_info = DrugInfo(query_inn="тест", matched_inn="тест", trade_names="Тестдраг")
    s2 = Stage2Result()
    s2.pk = PKParams(
        t_half_h=PKValue(value=3.0, unit="ч"),
        tmax_h=PKValue(value=1.0, unit="ч"),
    )

    inp = Stage3Input(drug_info=drug_info, s2=s2, test_drug_name="Генерик-Тест")
    result = generate_synopsis(inp, llm_fn=None)

    assert result.synopsis["n_total"] == "12+"
    assert "CVintra" in result.synopsis["sample_size_note"]
    assert result.synopsis["washout_days"] == "1"
    assert result.synopsis["vomit_criterion_h"] == "2.0"

def test_stage3_synopsis_no_llm_no_pk():
    """Stage 3 без LLM и без ФК данных: не падает."""
    from pipeline.stage3 import Stage3Input, generate_synopsis
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams

    drug_info = DrugInfo(query_inn="тест", matched_inn="тест")
    s2 = Stage2Result()
    s2.pk = PKParams()

    inp = Stage3Input(drug_info=drug_info, s2=s2, test_drug_name="Пустой-Тест")
    result = generate_synopsis(inp, llm_fn=None)

    assert result.synopsis["test_drug_name"] == "Пустой-Тест"
    assert result.synopsis["washout_days"] == "—"
    assert result.synopsis["vomit_criterion_h"] == "—"
    assert result.synopsis["t_half"] == "—"

def test_stage3_generate_synopsis_step_mock():
    """generate_synopsis_step с мок-LLM: возвращает структуру."""
    from pipeline.stage3 import (
        Stage3Input, compute_derived, collect_all_data,
        generate_synopsis_step, LLM_CALLS, _load_rule85,
    )
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    drug_info = DrugInfo(query_inn="тест", matched_inn="тест", trade_names="Реф")
    s2 = Stage2Result()
    s2.pk = PKParams(t_half_h=PKValue(value=4.0, unit="ч"), tmax_h=PKValue(value=1.5, unit="ч"))

    inp = Stage3Input(drug_info=drug_info, s2=s2, test_drug_name="Ген")
    computed = compute_derived(inp)
    all_data = collect_all_data(inp)
    rule85 = _load_rule85()

    def mock_llm(prompt):
        return '{"tasks": "Мок-задачи", "study_design": "Мок-дизайн"}'

    call_def = LLM_CALLS[0]
    result = generate_synopsis_step(call_def, inp, computed, all_data, rule85, mock_llm)

    assert result["call_id"] == "study_design_analysis"
    assert "tasks" in result["fields_received"]
    assert result["data"]["tasks"] == "Мок-задачи"
    assert result["prompt_len"] > 0
    assert result["response_len"] > 0

def test_stage3_full_with_llm():
    """Stage 3 с реальным LLM: все 19 секций заполняются (ибупрофен)."""
    from pipeline.config import DEEPSEEK_API_KEY
    if not DEEPSEEK_API_KEY:
        print("  SKIP: DEEPSEEK_API_KEY не задан")
        return
    from pipeline.stage3 import Stage3Input, generate_synopsis
    from pipeline.stage2 import find_pk_params
    from pipeline.stage1 import find_original

    drug = find_original("ибупрофен")
    s2 = find_pk_params(drug, use_llm=True)

    from openai import OpenAI
    from pipeline.config import DEEPSEEK_MODEL
    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

    def llm_fn(prompt):
        resp = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": "Ты эксперт по клиническим исследованиям. Отвечай валидным JSON."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2, max_tokens=8000,
        )
        return resp.choices[0].message.content or "{}"

    inp = Stage3Input(drug_info=drug, s2=s2, test_drug_name="Ибупрофен-Тест", strength="400 мг")
    result = generate_synopsis(inp, llm_fn=llm_fn)

    assert len(result.llm_calls_log) == 3, f"Ожидалось 3 LLM-вызова, получили {len(result.llm_calls_log)}"

    expected_fields = [
        "tasks", "study_design",
        "inclusion_criteria", "exclusion_criteria",
        "safety_analysis",
    ]
    for fld in expected_fields:
        val = result.synopsis.get(fld, "")
        if isinstance(val, list):
            assert len(val) >= 3, f"Секция {fld} список слишком короткий: {len(val)} элементов"
        else:
            assert val and len(str(val)) > 20, f"Секция {fld} пустая или слишком короткая: '{str(val)[:50]}'"

    assert result.synopsis["inn"] == "ибупрофен"
    assert result.synopsis["test_drug_name"] == "Ибупрофен-Тест"
    assert result.sources_used


def test_stage3_docx_import():
    """python-docx устанавливается корректно."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("test")
    assert len(doc.paragraphs) == 1

def test_stage3_generate_docx_full_sections():
    """generate_docx: работает со всеми секциями (включая списки)."""
    from pipeline.stage3 import Stage3Result, generate_docx

    result = Stage3Result(
        synopsis={
            "protocol_title": "Протокол исследования БЭ",
            "inn": "ибупрофен",
            "test_drug_name": "Ибупрофен-Тест",
            "reference_drug_name": "Нурофен",
            "study_objectives": "Оценить биоэквивалентность...",
            "study_design": "Рандомизированное открытое перекрёстное 2x2",
            "methodology": "Однократный приём натощак",
            "inclusion_criteria": ["Здоровые добровольцы 18-55 лет", "ИМТ 18.5-30 кг/м²", "Информированное согласие"],
            "exclusion_criteria": ["Гиперчувствительность к НПВС", "ЖКТ заболевания"],
            "pk_parameters": "AUC0-t, AUC0-∞, Cmax",
            "be_criteria": "90% ДИ для отношения в пределах 80.00–125.00%",
            "statistical_methods": "ANOVA на лог-трансформированных данных",
            "safety_analysis": "Мониторинг нежелательных явлений",
            "n_total": "28",
            "washout_days": "7",
            "timepoints_schedule": "0, 0.25, 0.5, 0.75, 1, 1.5, 2, 3, 4, 6, 8, 12, 24",
        },
        sources_used=[
            {"name": "Видаль", "url": "https://www.vidal.ru", "type": "drug"},
            {"name": "ОХЛП", "url": "local", "type": "regulatory"},
        ],
    )
    docx_bytes = generate_docx(result)
    assert len(docx_bytes) > 1000
    assert docx_bytes[:4] == b"PK\x03\x04"

def test_stage3_compute_derived_no_thalf():
    """compute_derived без T½: washout_days = None."""
    from pipeline.stage3 import Stage3Input, compute_derived
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    s2 = Stage2Result()
    s2.pk = PKParams(tmax_h=PKValue(value=2.0, unit="ч"))
    inp = Stage3Input(drug_info=DrugInfo(query_inn="x", matched_inn="x"), s2=s2)
    computed = compute_derived(inp)
    assert computed["washout_days"] is None
    assert computed["vomit_criterion_h"] == 4.0
    assert computed["t_half"] is None

def test_stage3_compute_derived_with_fda_flags():
    """compute_derived с FDA flags: NTI/HVD влияют на дизайн."""
    from pipeline.stage3 import Stage3Input, compute_derived
    from pipeline.stage2 import Stage2Result
    from pipeline.models import DrugInfo, PKParams, PKValue

    s2 = Stage2Result()
    s2.pk = PKParams(t_half_h=PKValue(value=6.0, unit="ч"), tmax_h=PKValue(value=2.0, unit="ч"), cvintra_pct=PKValue(value=35.0, unit="%"))
    s2.fda_psg_result = {"is_nti": True, "is_hvd": False, "is_replicated": False}
    inp = Stage3Input(drug_info=DrugInfo(query_inn="x", matched_inn="x"), s2=s2)
    computed = compute_derived(inp)
    assert computed["is_nti"] is True
    design = computed.get("design", {})
    assert "NTI" in design.get("rationale", "") or "90" in design.get("be_limits", "")


# ═══════════════════════════════════════════════
# Runner
# ═══════════════════════════════════════════════

if __name__ == "__main__":
    tests = [(k, v) for k, v in sorted(globals().items()) if k.startswith("test_")]

    passed = failed = skipped = 0
    groups = {
        "Stage 1":          [t for t in tests if "stage1" in t[0]],
        "Видаль вещество":  [t for t in tests if "vidal_mol" in t[0]],
        "Видаль препарат":  [t for t in tests if "vidal_drug" in t[0]],
        "e-Drug3D":         [t for t in tests if "edrug3d" in t[0]],
        "OSP":              [t for t in tests if t[0].startswith("test_osp")],
        "CVintra PMC":      [t for t in tests if "cvintra_pmc" in t[0]],
        "DrugBank":         [t for t in tests if "drugbank" in t[0]],
        "ОХЛП":             [t for t in tests if "ohlp" in t[0]],
        "FDA PSG":          [t for t in tests if "fda_psg" in t[0]],
        "Stage 2 pipeline": [t for t in tests if "stage2" in t[0]],
        "Data integrity":   [t for t in tests if t[0].startswith("test_data")],
        "Stage 3":          [t for t in tests if "stage3" in t[0] or "sample_size" in t[0] or "timepoints" in t[0] or "determine_design" in t[0]],
        "LLM":              [t for t in tests if "with_llm" in t[0] or "source_tag" in t[0]],
    }

    for group, gtests in groups.items():
        if not gtests:
            continue
        print(f"\n── {group} ──")
        for name, fn in gtests:
            import io
            from contextlib import redirect_stdout
            buf = io.StringIO()
            try:
                with redirect_stdout(buf):
                    fn()
                out = buf.getvalue().strip()
                if "SKIP" in out:
                    print(f"  SKIP  {name}")
                    skipped += 1
                else:
                    print(f"  PASS  {name}")
                    passed += 1
            except AssertionError as e:
                print(f"  FAIL  {name}: {e}")
                failed += 1
            except Exception as e:
                print(f"  ERR   {name}: {type(e).__name__}: {e}")
                failed += 1

    print(f"\n{'='*55}")
    print(f"  {passed} passed | {failed} failed | {skipped} skipped | {passed+failed+skipped} total")
    if failed:
        sys.exit(1)

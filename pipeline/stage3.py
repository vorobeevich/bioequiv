"""
Стадия 3 — Генерация синопсиса протокола исследования биоэквивалентности.

Архитектура (v3):
  ~20 полей генерируются ПРОГРАММНО (шаблоны + формулы из instructions.docx)
  3 LLM-вызова — там, где нужен анализ данных препарата:

  CALL 1: study_design_analysis
    → tasks, study_design
    ← ОХЛП (dosing, pk, form), Видаль, FDA PSG
    LLM определяет кратность дозирования (одн./многокр.) по данным ОХЛП
    и обосновывает выбранный дизайн исследования.

  CALL 2: criteria
    → inclusion_criteria, exclusion_criteria, withdrawal_criteria
    ← ОХЛП (contra, precautions, pregnancy, interactions), Видаль, Rule 85

  CALL 3: drug_safety
    → test_drug_details, reference_drug_details, safety_analysis
    ← ОХЛП (composition, form, excipients, storage, adverse, overdose), Видаль, DrugBank

Вход: результаты Stage 1 + Stage 2 + пользовательские параметры
Выход: заполненный синопсис (dict полей) + Word-документ (.docx)
"""

import math
import os
import io
import json
import datetime
from dataclasses import dataclass, field
from typing import Optional, Dict, Any, List, Callable

from .sample_size import calc_sample_size, determine_design
from .timepoints import generate_timepoints
from .stage2 import Stage2Result
from .models import DrugInfo

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RULE85_PATH = os.path.join(BASE_DIR, "docs", "rule85_context.md")

SOURCES = {
    "vidal": "https://www.vidal.ru",
    "drugbank": "https://go.drugbank.com",
    "fda_psg": "https://www.accessdata.fda.gov/scripts/cder/psg/index.cfm",
    "ohlp": "https://lk.regmed.ru/Register/EAEU_SmPC",
    "cvintra_pmc": "https://pmc.ncbi.nlm.nih.gov/articles/PMC6989220/",
    "eaeu": "https://portal.eaeunion.org/sites/commonprocesses/ru-ru/Pages/DrugRegistrationDetails.aspx",
    "edrug3d": "https://chemoinfo.ipmc.cnrs.fr/TMP/tmp.81675/e-Drug3D_2162_PK.txt",
    "osp": "https://www.open-systems-pharmacology.org/",
    "rule85": "Решение Совета ЕАЭК от 03.11.2016 N 85",
}

_MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}

# ─────────────────────────────────────────────────────────────
# Data classes
# ─────────────────────────────────────────────────────────────

@dataclass
class Stage3Input:
    drug_info: DrugInfo
    s2: Stage2Result
    test_drug_name: str = ""
    sponsor: str = ""
    research_center: str = ""
    bioanalytical_lab: str = ""
    dosage_form: str = ""
    strength: str = ""
    # New v2 parameters
    fasting_fed: str = ""              # "fasting" | "fed" | "both" | ""
    cv_intra_user: float = 0.0         # user-provided CVintra, 0 = auto
    use_rsabe: bool = False            # explicit RSABE requirement
    design_preference: str = ""        # "2x2" | "replicated" | "parallel" | ""
    study_phases: str = "single"       # "single" | "multiple"
    gender: str = ""                   # "both" | "male" | "female" | ""
    age_range: str = ""                # e.g. "18-45"
    additional_requirements: str = ""  # free text


@dataclass
class Stage3Result:
    synopsis: Dict[str, str] = field(default_factory=dict)
    computed: Dict[str, Any] = field(default_factory=dict)
    sources_used: List[Dict[str, str]] = field(default_factory=list)
    llm_calls_log: List[Dict[str, Any]] = field(default_factory=list)


# ─────────────────────────────────────────────────────────────
# PK helpers
# ─────────────────────────────────────────────────────────────

def _get_pk_value(s2: Stage2Result, param_name: str) -> Optional[float]:
    pk = s2.pk
    if not pk:
        return None
    mapped = {"t_half": "t_half_h", "tmax": "tmax_h", "cv_intra": "cvintra_pct",
              "cmax": "cmax", "auc": "auc"}
    attr = mapped.get(param_name, param_name)
    val_obj = getattr(pk, attr, None)
    if val_obj is None:
        return None
    try:
        return float(val_obj.value)
    except (ValueError, TypeError, AttributeError):
        return None


def _get_cv_intra(s2: Stage2Result, cv_user: float = 0.0) -> Optional[float]:
    if cv_user and cv_user > 0:
        return cv_user
    cv = _get_pk_value(s2, "cv_intra")
    if cv:
        return cv
    if s2.fda_psg_result:
        thr = s2.fda_psg_result.get("cvintra_threshold")
        if thr:
            try:
                return float(thr)
            except (ValueError, TypeError):
                pass
    return None


def _load_rule85() -> str:
    try:
        with open(RULE85_PATH, encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return ""


# ─────────────────────────────────────────────────────────────
# RSABE calculations
# ─────────────────────────────────────────────────────────────

def calc_rsabe_limits(cv_pct: float) -> Optional[Dict[str, Any]]:
    """RSABE limits for Cmax when CVwR > 30%. Returns None if not applicable."""
    if cv_pct is None or cv_pct <= 30:
        return None
    cv = cv_pct / 100.0
    sw_r = math.sqrt(math.log(1 + cv ** 2))
    k = 0.760
    upper = math.exp(k * sw_r)
    lower = math.exp(-k * sw_r)
    return {
        "lower_pct": round(lower * 100, 2),
        "upper_pct": round(upper * 100, 2),
        "sw_r": round(sw_r, 4),
        "k": k,
        "cv_pct": cv_pct,
    }


# ─────────────────────────────────────────────────────────────
# Design details
# ─────────────────────────────────────────────────────────────

def _design_details(design_type: str) -> Dict[str, Any]:
    configs = {
        "2x2": {
            "n_periods": 2, "n_groups": 2, "ratio": "1:1",
            "sequences": ["TR", "RT"],
            "name_adj": "двухпериодное перекрёстное",
            "name_full": "открытого рандомизированного двухпериодного перекрёстного исследования в двух группах с приёмом однократной дозы",
            "n_washouts": 1,
        },
        "replicated": {
            "n_periods": 4, "n_groups": 2, "ratio": "1:1",
            "sequences": ["TRTR", "RTRT"],
            "name_adj": "четырёхпериодное реплицированное перекрёстное",
            "name_full": "открытого рандомизированного четырёхпериодного реплицированного перекрёстного исследования в двух группах",
            "n_washouts": 3,
        },
        "parallel": {
            "n_periods": 1, "n_groups": 2, "ratio": "1:1",
            "sequences": ["T", "R"],
            "name_adj": "в параллельных группах",
            "name_full": "открытого рандомизированного исследования в параллельных группах",
            "n_washouts": 0,
        },
    }
    return configs.get(design_type, configs["2x2"])


def _fasting_display(code: str) -> str:
    m = {"fasting": "натощак", "fed": "после приёма высококалорийной пищи",
         "both": "натощак и после приёма высококалорийной пищи"}
    return m.get(code, code if code else "натощак")


def _dose_adj(phases: str) -> str:
    return "многократного" if phases == "multiple" else "однократного"


def _dose_adj_short(phases: str) -> str:
    return "многократно" if phases == "multiple" else "однократно"


def _gender_display(code: str) -> str:
    m = {"male": "мужчины", "female": "женщины", "both": "мужчины и женщины"}
    return m.get(code, "мужчины и женщины")


def _date_ru() -> str:
    d = datetime.date.today()
    return f"{d.day} {_MONTHS_RU[d.month]} {d.year} года"


# ─────────────────────────────────────────────────────────────
# Compute derived values
# ─────────────────────────────────────────────────────────────

def compute_derived(inp: Stage3Input) -> Dict[str, Any]:
    s2 = inp.s2
    t_half = _get_pk_value(s2, "t_half")
    tmax = _get_pk_value(s2, "tmax")
    cv_intra = _get_cv_intra(s2, inp.cv_intra_user)

    washout_days = math.ceil(5 * t_half / 24) if t_half else None
    vomit_criterion_h = round(2 * tmax, 1) if tmax else None

    is_hvd = False
    is_nti = False
    is_replicated = False
    if s2.fda_psg_result:
        is_hvd = s2.fda_psg_result.get("is_hvd") in (True, "True", "true", "1")
        is_nti = s2.fda_psg_result.get("is_nti") in (True, "True", "true", "1")
        is_replicated = s2.fda_psg_result.get("is_replicated") in (True, "True", "true", "1")

    if inp.design_preference:
        pref_map = {"2x2": "2x2", "replicated": "replicated", "parallel": "parallel",
                     "реплицированный": "replicated", "параллельный": "parallel"}
        forced_design = pref_map.get(inp.design_preference, "")
    else:
        forced_design = ""

    if inp.use_rsabe and not forced_design:
        forced_design = "replicated"

    if forced_design:
        theta = 1.1111 if is_nti else 1.25
        if forced_design == "replicated":
            be_limits = "80.00–125.00% (расширение Cmax допускается при обосновании)"
        elif is_nti:
            be_limits = "90.00–111.11%"
        else:
            be_limits = "80.00–125.00%"
        design_info = {
            "design": forced_design, "theta": theta, "be_limits": be_limits,
            "rationale": f"Дизайн задан пользователем: {forced_design}.",
        }
    else:
        design_info = determine_design(cv_intra, is_hvd=is_hvd, is_nti=is_nti,
                                        is_replicated_fda=is_replicated)

    dd = _design_details(design_info["design"])

    use_rsabe = inp.use_rsabe or (
        design_info["design"] == "replicated" and cv_intra is not None and cv_intra > 30
    )
    rsabe = calc_rsabe_limits(cv_intra) if use_rsabe else None

    sample = None
    if cv_intra:
        sample = calc_sample_size(
            cv_intra, power=0.80, theta=design_info["theta"],
            design=design_info["design"], dropout_pct=15.0,
        )

    n_to_screen = None
    if sample:
        n_to_screen = math.ceil(sample["n_total"] / (1 - 0.20))
        if n_to_screen % 2 != 0:
            n_to_screen += 1

    timepoints = None
    if tmax and t_half:
        timepoints = generate_timepoints(tmax, t_half)

    # Fasting/fed
    if inp.fasting_fed:
        fasting_code = inp.fasting_fed
    elif s2.fda_psg_result:
        df = s2.fda_psg_result.get("design_fasting", "")
        dd_fed = s2.fda_psg_result.get("design_fed", "")
        if df and dd_fed:
            fasting_code = "both"
        elif dd_fed and not df:
            fasting_code = "fed"
        else:
            fasting_code = "fasting"
    else:
        fasting_code = "fasting"

    fasting_text = _fasting_display(fasting_code)

    # PK period duration (days of hospitalization per period)
    sampling_end_h = timepoints["end_time_h"] if timepoints else 24
    pk_period_days = math.ceil(sampling_end_h / 24) + 2

    # Study duration
    n_periods = dd["n_periods"]
    n_washouts = dd["n_washouts"]
    if washout_days and n_periods > 1:
        study_duration_days = 14 + n_periods * pk_period_days + n_washouts * washout_days + 7
    elif n_periods == 1:
        study_duration_days = 14 + pk_period_days + 7
    else:
        study_duration_days = None

    # Reference drug info
    drug_info = inp.drug_info
    inn = drug_info.matched_inn if drug_info else ""
    inn_latin = ""
    if s2.vidal_mol_result:
        inn_latin = s2.vidal_mol_result.get("name_latin", "")
    ref_drug = ""
    ref_holder = ""
    ref_country = ""
    if drug_info:
        tn = drug_info.trade_names or ""
        ref_drug = tn.split(",")[0].strip() if tn else ""
        hd = drug_info.holders or ""
        ref_holder = hd.split(",")[0].strip() if hd else ""
        cn = drug_info.countries or ""
        ref_country = cn.split(",")[0].strip() if cn else ""

    age_range = inp.age_range or "18-45"
    gender = inp.gender or "both"

    return {
        "t_half": t_half, "tmax": tmax, "cv_intra": cv_intra,
        "washout_days": washout_days, "vomit_criterion_h": vomit_criterion_h,
        "design": design_info, "design_details": dd,
        "sample_size": sample, "n_to_screen": n_to_screen,
        "timepoints": timepoints,
        "study_duration_days": study_duration_days,
        "pk_period_days": pk_period_days,
        "sampling_end_h": sampling_end_h,
        "fasting_code": fasting_code,
        "fasting_text": fasting_text,
        "fasting_or_fed": fasting_text,
        "dose_type": inp.study_phases,
        "dose_adj": _dose_adj(inp.study_phases),
        "dose_adj_short": _dose_adj_short(inp.study_phases),
        "gender": gender, "gender_text": _gender_display(gender),
        "age_range": age_range,
        "is_hvd": is_hvd, "is_nti": is_nti,
        "use_rsabe": use_rsabe, "rsabe": rsabe,
        "inn": inn, "inn_latin": inn_latin,
        "ref_drug_name": ref_drug, "ref_holder": ref_holder, "ref_country": ref_country,
    }


# ─────────────────────────────────────────────────────────────
# Data collection from Stage 2
# ─────────────────────────────────────────────────────────────

def collect_all_data(inp: Stage3Input) -> Dict[str, str]:
    s2 = inp.s2
    data = {}

    if s2.vidal_drug_result:
        vdr = s2.vidal_drug_result
        parts = []
        for fld in ["pharmacokinetics", "pharmacology", "form_details", "composition"]:
            v = vdr.get(fld, "")
            if v:
                parts.append(f"{fld}: {v}")
        if parts:
            data["vidal_drug"] = "\n\n".join(parts)

    if s2.vidal_mol_result:
        vm = s2.vidal_mol_result
        for fld in ["pharmacokinetics", "pharmacology", "indications", "contraindications"]:
            v = vm.get(fld, "")
            if v:
                data[f"vidal_mol_{fld}"] = v

    if s2.ohlp_result:
        ohlp = s2.ohlp_result
        for key in ["pk_text", "dosing_text", "contra_text", "interactions_text",
                     "adverse_text", "pd_text", "storage_text", "composition_text",
                     "form_text", "excipients_text", "indications_text",
                     "precautions_text", "pregnancy_text", "overdose_text",
                     "shelf_life_text"]:
            v = ohlp.get(key, "")
            if v and len(v) > 20:
                data[f"ohlp_{key}"] = v

    if s2.drugbank_result:
        db = s2.drugbank_result
        for fld in ["absorption", "half_life", "protein_binding",
                     "volume_of_distribution", "clearance", "metabolism",
                     "route_of_elimination"]:
            v = db.get(fld, "")
            if v:
                data[f"drugbank_{fld}"] = v

    if s2.fda_psg_result:
        fda = s2.fda_psg_result
        for fld in ["design_fasting", "design_fed", "strength", "subjects",
                     "analytes", "be_based_on", "waiver", "additional_comments",
                     "dissolution_info"]:
            v = fda.get(fld, "")
            if v:
                data[f"fda_psg_{fld}"] = v

    return data


# ─────────────────────────────────────────────────────────────
# PROGRAMMATIC TEMPLATE GENERATORS
# ─────────────────────────────────────────────────────────────

def _gen_protocol_id(inp: Stage3Input, computed: Dict) -> str:
    return f"{computed['inn']}-BE"


def _gen_protocol_title(inp: Stage3Input, computed: Dict) -> str:
    dd = computed["design_details"]
    sponsor = inp.sponsor or "ХХХХХ, Россия"
    ref = computed["ref_drug_name"]
    ref_h = computed["ref_holder"]
    fasting = computed["fasting_text"]
    test = inp.test_drug_name
    form = inp.dosage_form
    strength = inp.strength

    test_part = test
    if form and form.lower() not in test.lower():
        test_part += f", {form}"
    if strength and strength not in test:
        test_part += f", {strength}"

    ref_part = f"референтного препарата {ref}" if ref else "референтного препарата"
    if ref_h:
        ref_part += f" ({ref_h})"

    return (
        f"Открытое рандомизированное {dd['name_adj']} исследование "
        f"сравнительной фармакокинетики (биоэквивалентности) препаратов "
        f"{test_part} ({sponsor}) и {ref_part} {fasting} у здоровых добровольцев."
    )


def _gen_study_objectives(inp: Stage3Input, computed: Dict) -> str:
    test = inp.test_drug_name
    form = inp.dosage_form
    strength = inp.strength
    test_full = test
    if form:
        test_full += f", {form}"
    if strength:
        test_full += f", {strength}"

    ref = computed["ref_drug_name"]
    ref_h = computed["ref_holder"]
    ref_part = f"референтного препарата {ref}" if ref else "референтного препарата"
    if ref_h:
        ref_part += f" ({ref_h})"

    fasting = computed["fasting_text"]

    return (
        f"Основная цель:\n"
        f"Оценка сравнительной фармакокинетики и биоэквивалентности "
        f"препаратов {test_full} и {ref_part} {fasting} у здоровых добровольцев.\n\n"
        f"Дополнительная цель:\n"
        f"Сравнительная оценка безопасности {computed['dose_adj']} приёма "
        f"препаратов {test_full} и {ref_part} у здоровых добровольцев."
    )


def _gen_tasks(inp: Stage3Input, computed: Dict) -> str:
    inn = computed["inn"]
    test = inp.test_drug_name
    dose_adj = computed["dose_adj"]
    dose_adj_short = computed["dose_adj_short"]
    form = inp.dosage_form
    strength = inp.strength

    test_full = test
    if form:
        test_full += f", {form}"
    if strength:
        test_full += f", {strength}"

    return (
        f"1. Определить концентрацию {inn} в плазме крови добровольцев после "
        f"{dose_adj} применения сравниваемых препаратов {test_full}.\n"
        f"2. Оценить фармакокинетические параметры и относительную биодоступность "
        f"сравниваемых препаратов.\n"
        f"3. Оценить биоэквивалентность сравниваемых препаратов на основе "
        f"статистического анализа фармакокинетических данных.\n"
        f"4. Оценить профиль безопасности сравниваемых препаратов при {dose_adj_short} "
        f"применении (частоту возникновения нежелательных явлений (НЯ) / серьёзных "
        f"нежелательных явлений (СНЯ), изменения данных лабораторных исследований, "
        f"физикального осмотра, функций жизненно важных органов, показателей ЭКГ)."
    )


def _gen_study_design(inp: Stage3Input, computed: Dict) -> str:
    dd = computed["design_details"]
    design = computed["design"]
    test = inp.test_drug_name
    ref = computed["ref_drug_name"]
    t_half = computed["t_half"]
    washout = computed["washout_days"]

    washout_text = ""
    if washout and t_half:
        washout_text = (
            f"Периоды приёма препаратов исследования разделяются отмывочным периодом "
            f"продолжительностью не менее пяти периодов полувыведения (T½) "
            f"действующего вещества (≥ 5 × {t_half} ч = {washout} дней)."
        )

    return (
        f"Согласно действующим Правилам Евразийского экономического союза (ЕАЭС) "
        f"для сравнительной оценки фармакокинетических параметров исследуемого "
        f"препарата {test} и препарата {ref} стандартным дизайном исследования "
        f"биоэквивалентности является проведение {dd['name_full']}. "
        f"{washout_text}\n\n"
        f"{design.get('rationale', '')}"
    )


def _gen_methodology(inp: Stage3Input, computed: Dict) -> str:
    dd = computed["design_details"]
    n_per = dd["n_periods"]
    n_grp = dd["n_groups"]
    ratio = dd["ratio"]
    seqs = dd["sequences"]
    fasting = computed["fasting_text"]
    dose_short = computed["dose_adj_short"]
    sampling_end = computed["sampling_end_h"]
    pk_days = computed["pk_period_days"]
    washout = computed["washout_days"]
    sample = computed.get("sample_size")
    n_per_group = sample["n_per_group"] if sample else "N/2"
    tp = computed.get("timepoints")
    schedule = tp["schedule_text"] if tp else "определяется индивидуально"

    periods_list = ", ".join([f"Период {i+1}" for i in range(n_per)])

    # Group table
    group_lines = []
    for gi, seq in enumerate(seqs):
        details = []
        for pi, letter in enumerate(seq):
            drug_label = "исследуемый препарат (T)" if letter == "T" else "референтный препарат (R)"
            details.append(f"  Период {pi+1}: {drug_label}")
        group_lines.append(
            f"Группа {gi+1} (n = {n_per_group} доб.): Последовательность {seq}\n"
            + "\n".join(details)
        )
    group_table = "\n".join(group_lines)

    washout_section = ""
    if dd["n_washouts"] > 0 and washout:
        washout_section = (
            f"\nОтмывочный период между периодами ФК исследования составит "
            f"не менее {washout} дней (≥ 5 × T½) с момента приёма препарата."
        )

    fed_section = ""
    if computed["fasting_code"] in ("fed", "both"):
        fed_section = (
            "\n\nИсследование после приёма пищи: добровольцы получат стандартный "
            "высококалорийный завтрак (800–1000 ккал, ~50% жиров: белки ~150 ккал, "
            "углеводы ~250 ккал, жиры 500–600 ккал). Начало приёма пищи — за 30 минут "
            "до приёма препарата, завтрак должен быть съеден в течение 30 минут."
        )

    return (
        f"Настоящее исследование будет выполнено с участием здоровых добровольцев, "
        f"соответствующих критериям включения/невключения и подписавших "
        f"«Информационный листок добровольца с формой информированного согласия».\n\n"
        f"Исследование будет состоять из следующих периодов: периода скрининга, "
        f"{n_per} периодов ФК исследования ({periods_list})"
        f"{' с отмывочным периодом между ними' if dd['n_washouts'] > 0 else ''}"
        f" и периода последующего наблюдения.\n\n"
        f"Длительность периода скрининга составит от 1 до 14 дней, "
        f"длительность каждого периода ФК исследования составит {pk_days} дней.\n\n"
        f"Добровольцы будут распределены в соответствии с рандомизационным списком "
        f"в одну из {n_grp} групп в соотношении {ratio}:\n\n"
        f"{group_table}\n\n"
        f"Период скрининга: после подписания Информационного листка с формой "
        f"информированного согласия добровольцы пройдут процедуры скрининга для оценки "
        f"соответствия критериям включения/невключения.\n\n"
        f"Периоды фармакокинетического (ФК) исследования:\n"
        f"Добровольцы будут госпитализированы в исследовательский центр вечером "
        f"накануне приёма препаратов (не менее чем за 10 часов до приёма). "
        f"Утром в День 1 каждого Периода ФК исследования добровольцы получат "
        f"{dose_short} дозу исследуемого/референтного препарата {fasting}, "
        f"запивая 200 мл бутилированной негазированной воды комнатной температуры.\n\n"
        f"Добровольцы останутся в центре в течение как минимум {sampling_end:.0f} часов "
        f"после дозирования с целью отбора биообразцов для анализа фармакокинетики.\n\n"
        f"Отбор образцов крови будет проводиться в следующие временные точки: "
        f"{schedule}.\n\n"
        f"Ограничения: приём жидкости не допускается в течение 1 часа до и 2 часов "
        f"после приёма препарата (за исключением воды для запивания). Приём пищи "
        f"не допускается в течение 4 часов после приёма препарата. Стандартные приёмы "
        f"пищи предоставляются по расписанию исследовательского центра."
        f"{washout_section}{fed_section}\n\n"
        f"Период последующего наблюдения: через 7 дней после последнего приёма "
        f"препарата добровольцы посетят центр для оценки состояния здоровья."
    )


def _gen_sample_size_text(inp: Stage3Input, computed: Dict) -> str:
    sample = computed.get("sample_size")
    cv = computed.get("cv_intra")
    n_to_screen = computed.get("n_to_screen")

    if not sample or not cv:
        return (
            "CVintra не определён. Минимум по Правилу 85: 12 добровольцев, "
            "включённых в анализ."
        )

    return (
        f"При расчёте объёма выборки для данного исследования использовался "
        f"коэффициент внутрииндивидуальной вариабельности фармакокинетических "
        f"параметров (CVintra) = {cv:.1f}% для максимальной концентрации / "
        f"площади под кривой «концентрация—время» (Cmax/AUC₀₋ₜ).\n\n"
        f"Расчёт объёма выборки выполнен с помощью пакета PowerTOST "
        f"с использованием ПО «The R Project for Statistical Computing» "
        f"(https://www.r-project.org) версии не ниже 4.4.2.\n\n"
        f"{sample['formula_note']}\n\n"
        f"С учётом досрочного выбывания не более 15% включённых добровольцев, "
        f"в исследование будет включено {sample['n_total']} добровольцев.\n"
        f"С учётом возможного 20% отсева на скрининге в исследование будут "
        f"скринированы до {n_to_screen or '—'} добровольцев.\n"
        f"Добровольцы, досрочно завершившие исследование, не будут заменены."
    )


def _gen_study_periods(inp: Stage3Input, computed: Dict) -> str:
    dd = computed["design_details"]
    n_per = dd["n_periods"]
    washout = computed.get("washout_days")
    sampling_end = computed["sampling_end_h"]
    s_days = math.ceil(sampling_end / 24)

    lines = []
    lines.append(
        "Период скрининга (предварительное обследование добровольцев):\n"
        "Визит 1. (День -14 – День -1).\n"
        "Для оценки соответствия добровольца критериям включения/невключения "
        "должны быть известны все результаты клинических, лабораторных и иных "
        "обследований."
    )

    visit_num = 2
    day_cursor = 0
    for p in range(n_per):
        hosp_day = day_cursor
        drug_day = day_cursor + 1
        sample_end_day = drug_day + s_days
        discharge_day = sample_end_day

        lines.append(
            f"\nПериод {p+1} ФК исследования:\n"
            f"Визит {visit_num}. День {hosp_day} – День {discharge_day} (госпитализация)\n"
            f"  Госпитализация{'и рандомизация' if p == 0 else ''} — День {hosp_day}\n"
            f"  Приём препарата — День {drug_day}\n"
            f"  Отбор образцов крови — День {drug_day} – День {sample_end_day}\n"
            f"  Завершение госпитализации — День {discharge_day}"
        )
        visit_num += 1

        if p < n_per - 1 and washout:
            wash_end = drug_day + washout
            lines.append(
                f"\nОтмывочный период: День {discharge_day + 1} – День {wash_end} "
                f"({washout} дней от приёма препарата в Периоде {p+1})."
            )
            day_cursor = wash_end
        else:
            day_cursor = discharge_day

    followup_day = day_cursor + 7
    lines.append(
        f"\nПериод последующего наблюдения:\n"
        f"Визит {visit_num}. День {followup_day} (окно визита +2 дня)\n"
        f"Доброволец посетит центр через 7 дней с момента последнего приёма "
        f"препарата для оценки состояния здоровья."
    )

    lines.append(
        "\nНезапланированный визит:\n"
        "Проводится при необходимости. При наличии показаний может быть "
        "дополнительно выполнена любая из процедур исследования по решению "
        "Исследователя."
    )

    return "\n".join(lines)


def _gen_study_duration(inp: Stage3Input, computed: Dict) -> str:
    dur = computed.get("study_duration_days")
    pk_days = computed["pk_period_days"]
    washout = computed.get("washout_days")

    if dur is None:
        return "Продолжительность определяется индивидуально."

    parts = [
        f"Максимальная продолжительность участия в исследовании для одного "
        f"добровольца составит {dur} дней.",
        f"Период скрининга продлится от 1 до 14 дней.",
        f"Длительность Периодов ФК исследования составит по {pk_days} дня "
        f"(включая госпитализацию вечером накануне приёма препарата, "
        f"не менее чем за 10 часов до приёма препарата).",
    ]
    if washout:
        parts.append(
            f"Отмывочный период между Периодами ФК исследования продлится "
            f"{washout} дней с момента приёма препаратов исследования."
        )
    parts.append(
        "Визит последующего наблюдения проводится на 7-й день "
        "после последнего приёма препарата."
    )
    return " ".join(parts)


def _gen_pk_parameters(inp: Stage3Input, computed: Dict) -> str:
    inn = computed["inn"]
    return (
        f"В данном исследовании будет изучена фармакокинетика {inn} по данным "
        f"исходного соединения в плазме крови.\n\n"
        f"1. Следующие фармакокинетические параметры {inn} будут оцениваться "
        f"в качестве первичных:\n"
        f"  • Cmax — максимальная плазменная концентрация\n"
        f"  • AUC₀₋ₜ — площадь под кривой «плазменная концентрация — время» "
        f"с момента приёма до последней определяемой концентрации.\n\n"
        f"2. Следующие фармакокинетические параметры {inn} будут оцениваться "
        f"в качестве вторичных:\n"
        f"  • AUC₀₋∞ — площадь под кривой, экстраполированная до бесконечности\n"
        f"  • Tmax — время достижения максимальной концентрации\n"
        f"  • T½ — период полувыведения из плазмы\n"
        f"  • kel — константа скорости терминальной элиминации."
    )


def _gen_analytical_method(inp: Stage3Input, computed: Dict) -> str:
    inn = computed["inn"]
    return (
        f"Для определения концентрации аналита в плазме крови здоровых "
        f"добровольцев будет применяться метод высокоэффективной жидкостной "
        f"хроматографии с тандемным масс-селективным детектированием "
        f"(ВЭЖХ-МС/МС). Полная валидация биоаналитического метода определения "
        f"{inn} в плазме крови будет проведена в соответствии с рекомендациями "
        f"ЕАЭС и стандартными процедурами аналитической лаборатории."
    )


def _gen_be_criteria(inp: Stage3Input, computed: Dict) -> str:
    inn = computed["inn"]
    rsabe = computed.get("rsabe")

    base = (
        f"Вывод о биоэквивалентности сравниваемых препаратов будет сделан "
        f"с использованием подхода, основанного на оценке 90% доверительных "
        f"интервалов для отношений средних геометрических значений исследуемого "
        f"препарата (T) к препарату сравнения (R) для фармакокинетических "
        f"параметров AUC₀₋ₜ и Cmax {inn}."
    )

    if rsabe:
        return (
            f"{base}\n\n"
            f"Для AUC₀₋ₜ: препараты считаются биоэквивалентными, если границы "
            f"90% ДИ находятся в пределах 80,00–125,00% (α = 0.05).\n\n"
            f"Для Cmax: ввиду высокой внутрисубъектной вариабельности "
            f"(CVintra = {rsabe['cv_pct']:.1f}% > 30%) применяется "
            f"reference-scaled average bioequivalence (RSABE).\n"
            f"  swR = √ln(1 + CV²) = √ln(1 + {rsabe['cv_pct']/100:.4f}²) "
            f"= {rsabe['sw_r']:.4f}\n"
            f"  k = {rsabe['k']}\n"
            f"  Границы: exp(±k·swR) = exp(±{rsabe['k']} × {rsabe['sw_r']:.4f}) "
            f"= [{rsabe['lower_pct']:.2f}%; {rsabe['upper_pct']:.2f}%]\n\n"
            f"Решение ЕАЭК №85, п. 81–85."
        )

    return (
        f"{base}\n\n"
        f"Препараты считаются биоэквивалентными, если границы оценённых "
        f"доверительных интервалов для AUC₀₋ₜ и Cmax находятся в пределах "
        f"80,00–125,00% (α = 0.05) для изучаемого аналита.\n\n"
        f"Решение ЕАЭК №85, п. 81–85."
    )


def _gen_sample_size_calculation(inp: Stage3Input, computed: Dict) -> str:
    sample = computed.get("sample_size")
    cv = computed.get("cv_intra")
    if not sample or not cv:
        return (
            "CVintra не определён. Минимум по Правилу 85: 12 добровольцев."
        )

    cv_frac = cv / 100.0
    sigma2 = math.log(1 + cv_frac ** 2)
    delta = math.log(sample["theta_used"])

    return (
        f"Расчёт размера выборки:\n\n"
        f"Дизайн: {sample['design_used']}\n"
        f"CVintra = {cv:.1f}%\n"
        f"σ²_w = ln(1 + CV²) = ln(1 + {cv_frac:.4f}²) = {sigma2:.4f}\n"
        f"δ = ln(θ) = ln({sample['theta_used']}) = {delta:.4f}\n"
        f"α = {sample['alpha_used']}, мощность = {sample['power_used']*100:.0f}%\n\n"
        f"N = 2·(Z_α + Z_β)²·σ²_w / δ²\n\n"
        f"Минимум для анализа: {sample['n_evaluable']} добровольцев\n"
        f"С учётом 15% выбывания: {sample['n_total']} добровольцев "
        f"({sample['n_per_group']} на группу)\n\n"
        f"Метод расчёта: {sample['method']}\n"
        f"Основание: Решение ЕАЭК №85, п. 26, 81, 87."
    )


def _gen_statistical_methods(inp: Stage3Input, computed: Dict) -> str:
    inn = computed["inn"]
    rsabe = computed.get("rsabe")

    hypothesis = "H₀: GT/GR ≤ 0.80 или GT/GR ≥ 1.25"
    if rsabe:
        lower = rsabe["lower_pct"] / 100
        upper = rsabe["upper_pct"] / 100
        hypothesis = (
            f"Для AUC₀₋ₜ: H₀: GT/GR ≤ 0.80 или GT/GR ≥ 1.25\n"
            f"Для Cmax (RSABE): H₀: GT/GR ≤ {lower:.4f} или GT/GR ≥ {upper:.4f}"
        )

    rsabe_section = ""
    if rsabe:
        rsabe_section = (
            f"\n\nДля Cmax применяется reference-scaled approach (RSABE):\n"
            f"swR = √ln(1 + CV²) = {rsabe['sw_r']:.4f}, k = {rsabe['k']}\n"
            f"Границы: exp(±k·swR) = [{rsabe['lower_pct']:.2f}%; {rsabe['upper_pct']:.2f}%]"
        )

    return (
        f"Статистический анализ данных будет проведён в соответствии с подходами, "
        f"изложенными в Правилах проведения исследований биоэквивалентности "
        f"лекарственных препаратов в рамках ЕАЭС.\n\n"
        f"В исследовании будет проверяться гипотеза биоэквивалентности:\n"
        f"{hypothesis}\n\n"
        f"где GT и GR — средние геометрические ФК параметра {inn} исследуемого "
        f"препарата и препарата сравнения соответственно.\n\n"
        f"Для первичных и вторичных фармакокинетических параметров будут рассчитаны "
        f"описательные статистики (медиана, среднее арифметическое, среднее "
        f"геометрическое, минимальное и максимальное значение, стандартное "
        f"отклонение, коэффициент вариации).\n\n"
        f"Статистический анализ проводится в предположении о логнормальном "
        f"распределении параметров AUC₀₋ₜ и Cmax {inn}. "
        f"После логарифмического преобразования показатели анализируются с помощью "
        f"дисперсионного анализа (ANOVA).\n\n"
        f"Модель ANOVA включает фиксированные факторы:\n"
        f"  • Лекарственный препарат\n"
        f"  • Период\n"
        f"  • Последовательность приёма\n"
        f"  • Добровольцы, вложенные в последовательность\n\n"
        f"На основе результатов ANOVA рассчитываются 90% доверительные интервалы "
        f"для отношения средних геометрических (T/R)."
        f"{rsabe_section}\n\n"
        f"Решение ЕАЭК №85, п. 86–90."
    )


def _gen_blinding_randomization(inp: Stage3Input, computed: Dict) -> str:
    dd = computed["design_details"]
    n_grp = dd["n_groups"]
    ratio = dd["ratio"]
    seqs = dd["sequences"]
    sample = computed.get("sample_size")
    n_per_group = sample["n_per_group"] if sample else "N/2"

    seq_text = "/".join(seqs)

    return (
        f"Данное исследование является открытым. Однако для биоаналитической "
        f"лаборатории проводится заслепление: сотрудники лаборатории не будут "
        f"иметь доступа к рандомизационному списку до окончания биоаналитической "
        f"стадии исследования.\n\n"
        f"Добровольцы будут распределены в одну из {n_grp} групп в соответствии "
        f"с рандомизационным списком методом блочной рандомизации без стратификации "
        f"в соотношении {ratio} с использованием программы IWRS iRand.\n\n"
        f"Каждый доброволец будет рандомизирован с присвоением рандомизационного "
        f"номера, который определяет последовательность приёма исследуемого "
        f"препарата (T) и референтного препарата (R) — {seq_text}.\n\n"
        f"В каждую группу будет включено по {n_per_group} добровольцев. "
        f"В каждом периоде добровольцы получат исследуемый/референтный препарат "
        f"по схеме в соответствии с группой, в которую они распределены."
    )


def _gen_ethical_aspects(inp: Stage3Input) -> str:
    sponsor = inp.sponsor or "ХХХХХ"
    return (
        "Исследование будет проводиться согласно Протоколу, в строгом соответствии с:\n"
        "• Конституцией Российской Федерации;\n"
        "• Этическими принципами Хельсинкской декларации Всемирной медицинской "
        "ассоциации 1964 г. в последней редакции, принятой на 75-й Генеральной "
        "Ассамблее ВМА, Хельсинки, Финляндия, октябрь 2024 г;\n"
        "• Решением Совета ЕАЭК от 03.11.2016 № 79 «Об утверждении Правил "
        "Надлежащей клинической практики ЕАЭС»;\n"
        "• Решением Совета ЕАЭК от 03.11.2016 № 85 «Об утверждении Правил "
        "проведения исследований биоэквивалентности лекарственных препаратов "
        "в рамках ЕАЭС»;\n"
        "а также в соответствии с применимыми требованиями законодательства "
        "Российской Федерации.\n\n"
        f"Страхование жизни и здоровья добровольцев осуществляется компанией {sponsor}."
    )


def _gen_protocol_version() -> str:
    return f"1.0 от {_date_ru()}"


def generate_programmatic_fields(inp: Stage3Input, computed: Dict) -> Dict[str, str]:
    """Generate all programmatic (non-LLM) synopsis fields."""
    s = {}

    s["phase"] = "Исследование биоэквивалентности"
    s["inn"] = computed["inn"]
    s["inn_latin"] = computed["inn_latin"]
    s["test_drug_name"] = inp.test_drug_name
    s["sponsor"] = inp.sponsor or "ХХХХХ, Россия"
    s["research_center"] = inp.research_center
    s["bioanalytical_lab"] = inp.bioanalytical_lab
    s["dosage_form"] = inp.dosage_form
    s["strength"] = inp.strength
    s["reference_drug_name"] = computed["ref_drug_name"]
    s["reference_drug_holder"] = computed["ref_holder"]
    s["fasting_or_fed"] = computed["fasting_text"]

    design = computed.get("design", {})
    s["design_type"] = design.get("design", "2x2")
    s["be_limits"] = design.get("be_limits", "80.00–125.00%")
    s["design_rationale"] = design.get("rationale", "")

    sample = computed.get("sample_size")
    if sample:
        s["n_total"] = str(sample["n_total"])
        s["n_per_group"] = str(sample["n_per_group"])
        s["n_evaluable"] = str(sample["n_evaluable"])
        s["sample_size_note"] = sample["formula_note"]
    else:
        s["n_total"] = "12+"
        s["n_per_group"] = "6+"
        s["n_evaluable"] = "12"
        s["sample_size_note"] = "CVintra не определён. Минимум по Правилу 85: 12."

    s["n_to_screen"] = str(computed.get("n_to_screen") or "—")

    tp = computed.get("timepoints")
    if tp:
        s["timepoints_schedule"] = tp["schedule_text"]
        s["n_samples"] = str(tp["n_samples"])
        s["sampling_end_h"] = str(tp["end_time_h"])
        s["blood_total_2periods_ml"] = str(tp["total_blood_2periods_ml"])
        s["timepoints_rationale"] = tp["rationale"]

    wd = computed.get("washout_days")
    s["washout_days"] = str(wd) if wd is not None else "—"
    vc = computed.get("vomit_criterion_h")
    s["vomit_criterion_h"] = str(vc) if vc is not None else "—"
    s["t_half"] = f"{computed['t_half']} ч" if computed.get("t_half") else "—"
    s["tmax"] = f"{computed['tmax']} ч" if computed.get("tmax") else "—"
    s["cv_intra"] = f"{computed['cv_intra']}%" if computed.get("cv_intra") else "—"

    s["protocol_id"] = _gen_protocol_id(inp, computed)
    s["protocol_title"] = _gen_protocol_title(inp, computed)
    s["study_objectives"] = _gen_study_objectives(inp, computed)
    s["tasks"] = _gen_tasks(inp, computed)
    s["study_design"] = _gen_study_design(inp, computed)
    s["methodology"] = _gen_methodology(inp, computed)
    s["sample_size_text"] = _gen_sample_size_text(inp, computed)
    s["study_periods"] = _gen_study_periods(inp, computed)
    s["study_duration"] = _gen_study_duration(inp, computed)
    s["pk_parameters"] = _gen_pk_parameters(inp, computed)
    s["analytical_method"] = _gen_analytical_method(inp, computed)
    s["be_criteria"] = _gen_be_criteria(inp, computed)
    s["sample_size_calculation"] = _gen_sample_size_calculation(inp, computed)
    s["statistical_methods"] = _gen_statistical_methods(inp, computed)
    s["blinding_randomization"] = _gen_blinding_randomization(inp, computed)
    s["ethical_aspects"] = _gen_ethical_aspects(inp)
    s["protocol_version"] = _gen_protocol_version()

    return s


# ─────────────────────────────────────────────────────────────
# LLM CALLS (2 instead of 6)
# ─────────────────────────────────────────────────────────────

LLM_CALLS = [
    {
        "id": "study_design_analysis",
        "name": "Анализ дизайна и задач исследования",
        "fields": ["tasks", "study_design"],
        "data_keys": [
            "ohlp_dosing_text", "ohlp_pk_text", "ohlp_form_text",
            "vidal_drug", "vidal_mol_pharmacokinetics",
            "fda_psg_design_fasting", "fda_psg_design_fed",
            "fda_psg_additional_comments", "fda_psg_strength",
        ],
        "description": "Определение кратности дозирования и обоснование дизайна",
    },
    {
        "id": "criteria",
        "name": "Критерии включения/невключения/исключения",
        "fields": ["inclusion_criteria", "exclusion_criteria", "withdrawal_criteria"],
        "data_keys": [
            "ohlp_contra_text", "ohlp_precautions_text", "ohlp_pregnancy_text",
            "ohlp_interactions_text", "ohlp_adverse_text", "ohlp_indications_text",
            "vidal_mol_contraindications", "vidal_mol_indications",
        ],
        "description": "Критерии включения, невключения и досрочного выбывания",
    },
    {
        "id": "drug_safety",
        "name": "Описание препаратов и безопасность",
        "fields": ["test_drug_details", "reference_drug_details", "safety_analysis"],
        "data_keys": [
            "vidal_drug", "ohlp_composition_text", "ohlp_form_text",
            "ohlp_excipients_text", "ohlp_storage_text", "ohlp_shelf_life_text",
            "ohlp_adverse_text", "ohlp_overdose_text", "ohlp_precautions_text",
            "drugbank_metabolism", "drugbank_route_of_elimination",
            "fda_psg_strength",
        ],
        "description": "Описание исследуемого и референтного препаратов, безопасность",
    },
]


def _common_context(inp: Stage3Input, computed: Dict[str, Any]) -> str:
    design = computed.get("design", {})
    sample = computed.get("sample_size")
    tp = computed.get("timepoints")

    return (
        f"МНН: {computed['inn']} ({computed['inn_latin']})\n"
        f"Референтный препарат: {computed['ref_drug_name']} ({computed['ref_holder']})\n"
        f"Исследуемый препарат (генерик): {inp.test_drug_name}\n"
        f"Лекарственная форма: {inp.dosage_form}\n"
        f"Дозировка: {inp.strength}\n"
        f"Условия приёма: {computed['fasting_text']}\n"
        f"Пол добровольцев: {computed['gender_text']}\n"
        f"Возраст: {computed['age_range']} лет\n"
        f"Применение: {computed['dose_adj']}\n\n"
        f"Дизайн: {design.get('design', '2x2')} — {design.get('rationale', '')}\n"
        f"Границы БЭ: {design.get('be_limits', '80.00–125.00%')}\n"
        f"NTI: {'да' if computed.get('is_nti') else 'нет'}, "
        f"HVD: {'да' if computed.get('is_hvd') else 'нет'}\n"
        f"T½ = {computed.get('t_half', '?')} ч, "
        f"Tmax = {computed.get('tmax', '?')} ч, "
        f"CVintra = {computed.get('cv_intra', '?')}%\n"
        f"Отмывочный период: {computed.get('washout_days', '?')} дней, "
        f"Критерий рвоты: {computed.get('vomit_criterion_h', '?')} ч\n"
        f"Выборка: {sample.get('n_total', '?') if sample else '?'} доб.\n"
        f"График крови: {tp.get('schedule_text', '?') if tp else '?'}"
    )


def _build_prompt_for_call(
    call_def: dict,
    inp: Stage3Input,
    computed: Dict[str, Any],
    all_data: Dict[str, str],
    rule85: str,
) -> str:
    ctx = _common_context(inp, computed)

    relevant_data = ""
    for key in call_def["data_keys"]:
        if key in all_data:
            txt = all_data[key]
            if len(txt) > 3000:
                txt = txt[:3000] + "..."
            relevant_data += f"\n### [{key}]\n{txt}\n"

    call_id = call_def["id"]

    if call_id == "study_design_analysis":
        rule85_section = _extract_rule85_sections(
            rule85, ["Дизайн", "однократн", "многократн", "перекрёстн"]
        )
        dd = computed["design_details"]
        design = computed["design"]
        t_half = computed["t_half"]
        tmax = computed.get("tmax")
        cv = computed.get("cv_intra")
        washout = computed.get("washout_days")
        test = inp.test_drug_name
        ref = computed["ref_drug_name"]
        fasting = computed["fasting_text"]
        dose_adj = computed["dose_adj"]
        inn = computed["inn"]

        programmatic_tasks = _gen_tasks(inp, computed)
        programmatic_design = _gen_study_design(inp, computed)

        template_instructions = (
            f"У тебя две задачи:\n\n"
            f"### Задача 1: Поле «tasks» — Задачи исследования\n"
            f"Текущий автоматический текст задач:\n"
            f"---\n{programmatic_tasks}\n---\n\n"
            f"Улучши этот текст, добавив клинически обоснованные детали на основе "
            f"данных ОХЛП/Видаль для конкретного препарата {inn}.\n\n"
            f"ВАЖНО — кратность дозирования:\n"
            f"Текущая настройка: {dose_adj} применение.\n"
            f"Проанализируй данные ОХЛП (раздел «Способ применения и дозы», "
            f"«Фармакокинетика», лекарственная форма) и определи:\n"
            f"- Нужно ли МНОГОКРАТНОЕ дозирование? Оно применяется ТОЛЬКО если:\n"
            f"  1) Однократная доза плохо переносится здоровыми добровольцами, ИЛИ\n"
            f"  2) Невозможно зафиксировать концентрацию после однократного приёма "
            f"(очень низкие концентрации в плазме), ИЛИ\n"
            f"  3) Лекарственная форма модифицированного высвобождения требует "
            f"достижения steady-state.\n"
            f"- Если НИ ОДНО из условий не выполняется → оставь ОДНОКРАТНОЕ.\n"
            f"- Если определил многократное — укажи причину в тексте задач.\n\n"
            f"### Задача 2: Поле «study_design» — Дизайн исследования\n"
            f"Текущий автоматический текст дизайна:\n"
            f"---\n{programmatic_design}\n---\n\n"
            f"Улучши обоснование дизайна, используя:\n"
            f"- Данные FDA PSG (если есть) — конкретные рекомендации по дизайну\n"
            f"- Данные ОХЛП о фармакокинетике препарата\n"
            f"- Клиническое обоснование выбранного дизайна ({dd['name_adj']})\n"
            f"- Почему именно этот дизайн (а не другой) оптимален для данного препарата\n"
            f"- T½ = {t_half} ч → отмывочный период ≥ {washout} дней\n"
            f"- CVintra = {cv}% → {'высокая вариабельность, реплицированный дизайн' if cv and cv >= 30 else 'стандартная вариабельность'}\n\n"
            f"Сохрани всю фактическую информацию из автоматического текста, "
            f"но дополни клиническим обоснованием."
        )

    elif call_id == "criteria":
        rule85_section = _extract_rule85_sections(
            rule85, ["Субъекты", "Критерии исключения"]
        )
        vomit = computed.get("vomit_criterion_h")
        tmax = computed.get("tmax")
        gender = computed["gender_text"]
        age = computed["age_range"]

        template_instructions = (
            f"Сгенерируй 3 секции критериев. Используй шаблоны ниже, адаптируя "
            f"к конкретному препарату.\n\n"
            f"ВАЖНО:\n"
            f"- Пол добровольцев: {gender}. Если препарат только для мужчин, "
            f"убери упоминания женщин. Если только для женщин — убери мужчин.\n"
            f"- Возраст: {age} лет.\n"
            f"- Если препарат противопоказан беременным → добавь тест на беременность "
            f"в критерии исключения.\n"
            f"- Критерий рвоты: {vomit} ч (удвоенная величина Tmax = {tmax} ч).\n"
            f"- Используй конкретные противопоказания из данных ОХЛП/Видаль.\n"
        )

    elif call_id == "drug_safety":
        rule85_section = ""
        fasting = computed["fasting_text"]
        dose_short = computed["dose_adj_short"]
        ref_drug = computed["ref_drug_name"]
        ref_holder = computed["ref_holder"]

        template_instructions = (
            f"Сгенерируй 3 секции.\n\n"
            f"test_drug_details — описание исследуемого препарата:\n"
            f"  Лекарственная форма: {inp.dosage_form}\n"
            f"  Дозировка: {inp.strength}\n"
            f"  Состав, вспомогательные вещества: из данных ОХЛП/Видаль\n"
            f"  Схема приёма: «Каждый доброволец примет {dose_short} {fasting} "
            f"по ... препарата ... в последовательности по рандомизации»\n"
            f"  Условия хранения: из ОХЛП\n\n"
            f"reference_drug_details — описание референтного препарата ({ref_drug}):\n"
            f"  МНН: {computed['inn']}\n"
            f"  Та же структура что и для исследуемого. Данные из ОХЛП/Видаль.\n"
            f"  Держатель РУ: {ref_holder}\n"
            f"  Укажи: «Препарат {ref_drug} является оригинальным препаратом, "
            f"зарегистрированным на территории РФ.»\n"
            f"  «Сравниваемые препараты содержат одно и то же действующее вещество, "
            f"выпускаются в одинаковой лекарственной форме.»\n\n"
            f"safety_analysis — анализ безопасности:\n"
            f"  Базовый текст + специфичные для препарата анализы из ОХЛП "
            f"(нежелательные реакции, передозировка)."
        )
    else:
        rule85_section = ""
        template_instructions = ""

    fields_desc = {
        "tasks": (
            "Задачи исследования (нумерованный список, 4-5 пунктов). "
            "Учти кратность дозирования (однократное/многократное) "
            "на основе анализа данных ОХЛП."
        ),
        "study_design": (
            "Дизайн исследования: обоснование выбранного дизайна "
            "с учётом специфики препарата, отмывочный период, "
            "клиническое обоснование. Развёрнутый текст на 2-3 абзаца."
        ),
        "inclusion_criteria": (
            "Критерии включения (нумерованный список, ~10-15 пунктов, "
            "на основе шаблона + специфика препарата из ОХЛП)"
        ),
        "exclusion_criteria": (
            "Критерии невключения (нумерованный список, ~15-20 пунктов, "
            "включая специфичные для этого препарата из противопоказаний)"
        ),
        "withdrawal_criteria": (
            "Критерии досрочного выбывания (нумерованный список, "
            "включая критерий рвоты и при необходимости тест на беременность)"
        ),
        "test_drug_details": (
            "Описание исследуемого препарата (форма, состав, дозировка, "
            "вспомогательные вещества, схема приёма, хранение, производитель)"
        ),
        "reference_drug_details": (
            "Описание референтного препарата (форма, состав, дозировка, "
            "вспомогательные вещества, схема приёма, хранение, производитель, "
            "номер РУ, обоснование выбора)"
        ),
        "safety_analysis": (
            "Анализ параметров безопасности (НЯ/СНЯ, лабораторные данные, "
            "ЭКГ, физикальный осмотр, витальные показатели)"
        ),
    }

    fields_block = "\n".join(
        f'{i+1}. "{f}" — {fields_desc.get(f, f)}'
        for i, f in enumerate(call_def["fields"])
    )

    prompt = (
        f"Ты — эксперт по клиническим исследованиям биоэквивалентности.\n\n"
        f"## Контекст исследования\n{ctx}\n\n"
        f"## Инструкции\n{template_instructions}\n\n"
        f"## Данные из источников\n"
        f"{relevant_data if relevant_data.strip() else '(нет дополнительных данных)'}\n\n"
    )
    if rule85_section:
        prompt += f"## Нормативная база (Решение ЕАЭК №85)\n{rule85_section}\n\n"

    if inp.additional_requirements:
        prompt += f"## Дополнительные требования заказчика\n{inp.additional_requirements}\n\n"

    prompt += (
        f"## Задание\n"
        f"Сгенерируй JSON с полями (все значения — строки на русском языке):\n\n"
        f"{fields_block}\n\n"
        f"Требования:\n"
        f"- Пиши подробно, профессионально, на русском языке\n"
        f"- Ссылайся на пункты Решения ЕАЭК №85 где уместно\n"
        f"- Используй данные из источников для специфики именно этого препарата\n"
        f"- Отвечай ТОЛЬКО валидным JSON объектом без markdown-обёрток\n"
    )
    return prompt


def _extract_rule85_sections(rule85: str, keywords: List[str]) -> str:
    if not rule85:
        return ""
    sections = rule85.split("\n## ")
    relevant = []
    for section in sections:
        for kw in keywords:
            if kw.lower() in section.lower():
                relevant.append("## " + section if not section.startswith("#") else section)
                break
    return "\n\n".join(relevant)[:4000]


def _parse_llm_json(raw: str) -> dict:
    raw = raw.strip()
    if raw.startswith("```"):
        lines = raw.split("\n")
        raw = "\n".join(lines[1:])
        if raw.endswith("```"):
            raw = raw[:-3]
        raw = raw.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            try:
                return json.loads(raw[start:end+1])
            except json.JSONDecodeError:
                pass
    return {}


def generate_synopsis_step(
    call_def: dict,
    inp: Stage3Input,
    computed: Dict[str, Any],
    all_data: Dict[str, str],
    rule85: str,
    llm_fn: Callable,
) -> Dict[str, Any]:
    prompt = _build_prompt_for_call(call_def, inp, computed, all_data, rule85)
    raw = llm_fn(prompt)
    parsed = _parse_llm_json(raw)
    return {
        "call_id": call_def["id"],
        "call_name": call_def["name"],
        "fields_requested": call_def["fields"],
        "fields_received": list(parsed.keys()),
        "data": parsed,
        "prompt_len": len(prompt),
        "response_len": len(raw),
    }


# ─────────────────────────────────────────────────────────────
# ORCHESTRATOR
# ─────────────────────────────────────────────────────────────

def generate_synopsis(
    inp: Stage3Input, llm_fn=None, progress_fn=None
) -> Stage3Result:
    computed = compute_derived(inp)
    all_data = collect_all_data(inp)
    rule85 = _load_rule85()

    synopsis = generate_programmatic_fields(inp, computed)

    sources_used = []
    llm_calls_log = []
    _collect_source_links(inp.s2, sources_used)

    if llm_fn:
        for call_def in LLM_CALLS:
            if progress_fn:
                progress_fn(call_def["name"], "started")
            try:
                result = generate_synopsis_step(
                    call_def, inp, computed, all_data, rule85, llm_fn
                )
                synopsis.update(result["data"])
                llm_calls_log.append(result)
                if progress_fn:
                    progress_fn(call_def["name"], "done", result)
            except Exception as e:
                err_info = {"call_id": call_def["id"], "error": str(e)}
                llm_calls_log.append(err_info)
                if progress_fn:
                    progress_fn(call_def["name"], "error", err_info)

    return Stage3Result(
        synopsis=synopsis,
        computed=computed,
        sources_used=sources_used,
        llm_calls_log=llm_calls_log,
    )


# ─────────────────────────────────────────────────────────────
# SOURCE LINKS
# ─────────────────────────────────────────────────────────────

def _collect_source_links(s2: Stage2Result, sources: List[Dict[str, str]]):
    sources.append({"name": "Решение ЕАЭК №85", "url": SOURCES["rule85"], "type": "regulation"})
    sources.append({"name": "Реестр ЕАЭС", "url": SOURCES["eaeu"], "type": "registry"})

    if s2.vidal_drug_result:
        url = s2.vidal_drug_result.get("drug_url", "") or SOURCES["vidal"]
        sources.append({"name": f"Видаль: {s2.vidal_drug_result.get('drug_name', '')}", "url": url, "type": "drug"})
    if s2.vidal_mol_result:
        url = s2.vidal_mol_result.get("url", "") or SOURCES["vidal"]
        sources.append({"name": f"Видаль: {s2.vidal_mol_result.get('name_ru', '')}", "url": url, "type": "molecule"})
    if s2.drugbank_result:
        url = s2.drugbank_result.get("url", "") or SOURCES["drugbank"]
        sources.append({"name": f"DrugBank: {s2.drugbank_result.get('matched_name', '')}", "url": url, "type": "molecule"})
    if s2.ohlp_result:
        sources.append({"name": "ОХЛП (ГРЛС)", "url": SOURCES["ohlp"], "type": "drug"})
    if s2.fda_psg_result:
        url = s2.fda_psg_result.get("pdf_url", "") or SOURCES["fda_psg"]
        sources.append({"name": f"FDA PSG: {s2.fda_psg_result.get('substance', '')}", "url": url, "type": "guidance"})
    if s2.edrug3d_result:
        sources.append({"name": "e-Drug3D", "url": SOURCES["edrug3d"], "type": "database"})
    if s2.osp_result:
        sources.append({"name": "Open Systems Pharmacology", "url": SOURCES["osp"], "type": "database"})
    if s2.cvintra_pmc_result:
        sources.append({"name": "CVintra/PMC (Park et al. 2020)", "url": SOURCES["cvintra_pmc"], "type": "article"})


# ─────────────────────────────────────────────────────────────
# DOCX GENERATION
# ─────────────────────────────────────────────────────────────

def generate_docx(result: Stage3Result) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СИНОПСИС ПРОТОКОЛА")
    run.bold = True
    run.font.size = Pt(14)

    s = result.synopsis

    rows = [
        ("Название протокола:", s.get("protocol_title", "")),
        ("Идентификационный номер:", s.get("protocol_id", "")),
        ("Спонсор:", s.get("sponsor", "")),
        ("Исследовательский центр:", s.get("research_center", "")),
        ("Биоаналитическая лаборатория:", s.get("bioanalytical_lab", "")),
        ("Фаза:", s.get("phase", "")),
        ("Название исследуемого препарата:", s.get("test_drug_name", "")),
        ("Действующее вещество:", f"{s.get('inn', '')} ({s.get('inn_latin', '')})"),
        ("Цель исследования:", s.get("study_objectives", "")),
        ("Задачи исследования:", s.get("tasks", "")),
        ("Дизайн исследования:", s.get("study_design", "")),
        ("Методология исследования:", s.get("methodology", "")),
        ("Количество добровольцев:", s.get("sample_size_text", "")),
        ("Критерии включения:", s.get("inclusion_criteria", "")),
        ("Критерии невключения:", s.get("exclusion_criteria", "")),
        ("Критерии исключения:", s.get("withdrawal_criteria", "")),
        ("Исследуемый препарат (T):", s.get("test_drug_details", "")),
        ("Референтный препарат (R):", s.get("reference_drug_details", "")),
        ("Периоды исследования:", s.get("study_periods", "")),
        ("Продолжительность:", s.get("study_duration", "")),
        ("ФК параметры:", s.get("pk_parameters", "")),
        ("Аналитический метод:", s.get("analytical_method", "")),
        ("Критерии БЭ:", s.get("be_criteria", "")),
        ("Безопасность:", s.get("safety_analysis", "")),
        ("Расчёт выборки:", s.get("sample_size_calculation", "")),
        ("Стат. методы:", s.get("statistical_methods", "")),
        ("Рандомизация:", s.get("blinding_randomization", "")),
        ("Этика:", s.get("ethical_aspects", "")),
        ("Версия протокола:", s.get("protocol_version", "")),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        p_label = table.cell(i, 0).paragraphs[0]
        run_l = p_label.add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(10)
        p_value = table.cell(i, 1).paragraphs[0]
        run_v = p_value.add_run(str(value) if value else "")
        run_v.font.size = Pt(10)

    for cell in table.columns[0].cells:
        cell.width = Inches(2.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(4.0)

    if result.sources_used:
        doc.add_paragraph()
        p_src = doc.add_paragraph()
        p_src.add_run("Источники данных").bold = True
        for src in result.sources_used:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(src["name"])
            if src.get("url") and src["url"].startswith("http"):
                p.add_run(f" — {src['url']}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_objectives_text(s: dict) -> str:
    return s.get("study_objectives", "")
